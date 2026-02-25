"""Fill a Word template with term sheet data using find/replace.

This module takes an existing .docx template and fills in placeholders.
Placeholders use the format {{PLACEHOLDER_NAME}}.
"""

from io import BytesIO
from pathlib import Path

from docx import Document

from .models import TermSheet, InstrumentType, BoardRights


def _format_money(amount: float) -> str:
    """Format money without brackets or placeholders."""
    if amount >= 1_000_000_000:
        billions = amount / 1_000_000_000
        if billions == int(billions):
            return f"${int(billions)}B"
        return f"${billions:.1f}B"
    elif amount >= 1_000_000:
        millions = amount / 1_000_000
        if millions == int(millions):
            return f"${int(millions)}M"
        return f"${millions:.1f}M"
    elif amount >= 1_000:
        return f"${amount / 1_000:.0f}K"
    else:
        return f"${amount:,.0f}"


def _format_money_full(amount: float) -> str:
    """Format money as full dollar amount (e.g., $75,000)."""
    return f"${amount:,.0f}"


def _calculate_ownership(
    investment: float,
    post_money: float | None = None,
    pre_money: float | None = None,
) -> float:
    """Calculate ownership percentage."""
    if post_money:
        return (investment / post_money) * 100
    elif pre_money:
        return (investment / (pre_money + investment)) * 100
    return 0.0


def build_replacements(ts: TermSheet) -> dict[str, str]:
    """Build a dictionary of placeholder -> replacement value mappings."""
    replacements = {
        "{{COMPANY_NAME}}": ts.company_name,
        "{{COMPANY_NAME_UPPER}}": ts.company_name.upper(),
        "{{INVESTMENT_AMOUNT}}": _format_money(ts.investment_amount),
        "{{INVESTMENT_AMOUNT_FULL}}": _format_money_full(ts.investment_amount),
        "{{LEGAL_FEE_CAP}}": _format_money_full(ts.legal_fee_cap),
        "{{EXCLUSIVITY_DAYS}}": str(ts.exclusivity_days),
        "{{OPTION_POOL_PERCENT}}": f"{ts.option_pool_percent:.0f}%",
    }

    # Valuation
    if ts.post_money_valuation:
        replacements["{{VALUATION}}"] = _format_money(ts.post_money_valuation)
        replacements["{{VALUATION_TYPE}}"] = "post-money"
        ownership = _calculate_ownership(ts.investment_amount, post_money=ts.post_money_valuation)
        replacements["{{OWNERSHIP_PERCENT}}"] = f"{ownership:.1f}%"
    elif ts.pre_money_valuation:
        replacements["{{VALUATION}}"] = _format_money(ts.pre_money_valuation)
        replacements["{{VALUATION_TYPE}}"] = "pre-money"
        ownership = _calculate_ownership(ts.investment_amount, pre_money=ts.pre_money_valuation)
        replacements["{{OWNERSHIP_PERCENT}}"] = f"{ownership:.1f}%"
    elif ts.valuation_cap:
        replacements["{{VALUATION}}"] = _format_money(ts.valuation_cap)
        replacements["{{VALUATION_TYPE}}"] = "valuation cap"
        replacements["{{OWNERSHIP_PERCENT}}"] = "TBD at conversion"

    # Series
    if ts.series:
        replacements["{{SERIES}}"] = ts.series.upper()
        replacements["{{SERIES_LOWER}}"] = ts.series
    else:
        replacements["{{SERIES}}"] = "A"
        replacements["{{SERIES_LOWER}}"] = "a"

    # Instrument type
    if ts.instrument_type == InstrumentType.SAFE:
        replacements["{{INSTRUMENT_TYPE}}"] = "SAFE"
        replacements["{{INSTRUMENT_DESCRIPTION}}"] = "Simple Agreement for Future Equity (SAFE)"
    elif ts.instrument_type == InstrumentType.PRICED:
        series = ts.series or "A"
        replacements["{{INSTRUMENT_TYPE}}"] = f"Series {series} Preferred Stock"
        replacements["{{INSTRUMENT_DESCRIPTION}}"] = f"Series {series} Preferred Stock Financing"
    else:
        replacements["{{INSTRUMENT_TYPE}}"] = "Convertible Note"
        replacements["{{INSTRUMENT_DESCRIPTION}}"] = "Convertible Promissory Note"

    # Board rights
    if ts.board_rights == BoardRights.SEAT:
        replacements["{{BOARD_RIGHTS}}"] = "Board seat"
        replacements["{{BOARD_RIGHTS_TEXT}}"] = (
            "One director to be elected by the Series Preferred Stock and designated by Paradigm."
        )
    elif ts.board_rights == BoardRights.SEAT_AND_OBSERVER:
        replacements["{{BOARD_RIGHTS}}"] = "Board seat and observer rights"
        replacements["{{BOARD_RIGHTS_TEXT}}"] = (
            "One director to be elected by the Series Preferred Stock and designated by Paradigm. "
            "Company shall also invite a representative of Paradigm to attend all meetings of the "
            "Board in a nonvoting observer capacity."
        )
    elif ts.board_rights == BoardRights.OBSERVER:
        replacements["{{BOARD_RIGHTS}}"] = "Observer rights"
        replacements["{{BOARD_RIGHTS_TEXT}}"] = (
            "Company shall invite a representative of Paradigm to attend all meetings of the Board "
            "in a nonvoting observer capacity."
        )
    else:
        replacements["{{BOARD_RIGHTS}}"] = "None"
        replacements["{{BOARD_RIGHTS_TEXT}}"] = ""

    # Token rights
    if ts.token_rights.enabled:
        replacements["{{TOKEN_RIGHTS}}"] = "Yes"
        replacements["{{TOKEN_FLOOR}}"] = f"{ts.token_rights.token_floor_percent:.0f}%"
    else:
        replacements["{{TOKEN_RIGHTS}}"] = "No"
        replacements["{{TOKEN_FLOOR}}"] = "N/A"

    # Pro rata
    replacements["{{PRO_RATA}}"] = "Yes" if ts.pro_rata_rights else "No"

    # Names
    if ts.founder_name:
        replacements["{{FOUNDER_NAME}}"] = ts.founder_name
    if ts.dri_name:
        replacements["{{DRI_NAME}}"] = ts.dri_name

    # Custom terms
    if ts.custom_terms:
        replacements["{{CUSTOM_TERMS}}"] = ts.custom_terms
    else:
        replacements["{{CUSTOM_TERMS}}"] = ""

    return replacements


def replace_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    """Replace placeholders in a paragraph while preserving formatting."""
    # Build full text from all runs
    full_text = "".join(run.text for run in paragraph.runs)

    # Check if any replacements are needed
    if not any(key in full_text for key in replacements):
        return

    # Perform replacements
    new_text = full_text
    for placeholder, value in replacements.items():
        new_text = new_text.replace(placeholder, value)

    # If text changed, update the runs
    if new_text != full_text:
        # Clear all runs except first, put all text in first run
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""


def fill_template(template_path: str | Path, ts: TermSheet) -> bytes:
    """Fill a Word template with term sheet data.

    Args:
        template_path: Path to the .docx template file
        ts: TermSheet data to fill in

    Returns:
        bytes: The filled document as bytes
    """
    template_path = Path(template_path)
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(template_path)
    replacements = build_replacements(ts)

    # Replace in all paragraphs
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacements)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, replacements)

    # Replace in headers and footers
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_in_paragraph(paragraph, replacements)
        for paragraph in section.footer.paragraphs:
            replace_in_paragraph(paragraph, replacements)

    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
