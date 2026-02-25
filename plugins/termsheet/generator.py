"""Term sheet document generation.

NOTE: For strict template compliance, use generate_term_sheet_docx_strict()
from the template module instead of generate_term_sheet_docx().
"""

from io import BytesIO

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .models import TermSheet, InstrumentType, BoardRights
from .template import generate_term_sheet_docx_strict  # noqa: F401 — public re-export


def _format_money(amount: float) -> str:
    if amount >= 1_000_000_000:
        return f"${amount / 1_000_000_000:.1f}B"
    elif amount >= 1_000_000:
        return f"${amount / 1_000_000:.1f}M"
    elif amount >= 1_000:
        return f"${amount / 1_000:.0f}K"
    else:
        return f"${amount:,.0f}"


def _calculate_ownership(
    investment: float,
    post_money: float | None = None,
    pre_money: float | None = None,
) -> float:
    if post_money:
        return (investment / post_money) * 100
    elif pre_money:
        return (investment / (pre_money + investment)) * 100
    return 0.0


def generate_term_sheet_text(ts: TermSheet) -> str:
    lines = []

    lines.append(
        f"{ts.company_name.upper()} ",
    )

    if ts.instrument_type == InstrumentType.SAFE:
        lines.append("SIMPLE AGREEMENT FOR FUTURE EQUITY (SAFE)")
        lines.append("SUMMARY OF PROPOSED TERMS")
        lines.append("")
        lines.append("=" * 60)
        lines.append("")

        lines.append("INVESTMENT TERMS")
        lines.append("-" * 40)
        lines.append(f"Investment Amount:      {_format_money(ts.investment_amount)}")
        if ts.valuation_cap:
            lines.append(f"Valuation Cap:          {_format_money(ts.valuation_cap)}")
        if ts.discount_percent:
            lines.append(f"Discount:               {ts.discount_percent}%")
        lines.append("Instrument:             SAFE")

    elif ts.instrument_type == InstrumentType.PRICED:
        series = ts.series or "A"
        lines.append(f"SERIES {series.upper()} PREFERRED STOCK FINANCING")
        lines.append("SUMMARY OF PROPOSED TERMS")
        lines.append("")
        lines.append("=" * 60)
        lines.append("")

        lines.append("INVESTMENT & VALUATION")
        lines.append("-" * 40)
        lines.append(f"Investment Amount:      {_format_money(ts.investment_amount)}")
        if ts.post_money_valuation:
            lines.append(f"Post-Money Valuation:   {_format_money(ts.post_money_valuation)}")
            ownership = _calculate_ownership(
                ts.investment_amount, post_money=ts.post_money_valuation
            )
            lines.append(f"Ownership Post-Close:   {ownership:.1f}%")
        elif ts.pre_money_valuation:
            lines.append(f"Pre-Money Valuation:    {_format_money(ts.pre_money_valuation)}")
            ownership = _calculate_ownership(ts.investment_amount, pre_money=ts.pre_money_valuation)
            lines.append(f"Ownership Post-Close:   {ownership:.1f}%")

    elif ts.instrument_type == InstrumentType.CONVERTIBLE_NOTE:
        lines.append("CONVERTIBLE PROMISSORY NOTE")
        lines.append("SUMMARY OF PROPOSED TERMS")
        lines.append("")
        lines.append("=" * 60)
        lines.append("")

        lines.append("NOTE TERMS")
        lines.append("-" * 40)
        lines.append(f"Principal Amount:       {_format_money(ts.investment_amount)}")
        if ts.valuation_cap:
            lines.append(f"Valuation Cap:          {_format_money(ts.valuation_cap)}")
        if ts.discount_percent:
            lines.append(f"Discount:               {ts.discount_percent}%")

    lines.append("")
    lines.append(
        f"Option Pool:            {ts.option_pool_percent}% ({ts.option_pool_timing}-money)"
    )

    lines.append("")
    lines.append("SECURITIES")
    lines.append("-" * 40)
    if ts.instrument_type == InstrumentType.PRICED:
        series = ts.series or "A"
        lines.append(f"Security:               Series {series} Preferred Stock")
    lines.append(f"Liquidation Preference: {ts.liquidation_preference}")
    lines.append(f"Anti-Dilution:          {ts.anti_dilution}")

    lines.append("")
    lines.append("GOVERNANCE")
    lines.append("-" * 40)
    if ts.board_rights == BoardRights.SEAT:
        lines.append("Board Rights:           One director seat designated by Paradigm")
    elif ts.board_rights == BoardRights.SEAT_AND_OBSERVER:
        lines.append("Board Rights:           Board seat and observer rights")
    elif ts.board_rights == BoardRights.OBSERVER:
        lines.append("Board Rights:           Board observer rights")
    else:
        lines.append("Board Rights:           None")

    lines.append(
        f"Pro Rata Rights:        {'Yes - for Major Investors (Paradigm only)' if ts.pro_rata_rights else 'No'}"
    )

    lines.append("")
    lines.append("PROTECTIVE PROVISIONS")
    lines.append("-" * 40)
    lines.append("Consent of Paradigm required for:")
    lines.append("  • Incurrence of indebtedness > $1M")
    lines.append("  • Creation of new equity compensation plans")
    lines.append("  • Sale, license, or encumbrance of material IP")
    lines.append("  • Creation, sale, or issuance of any tokens")
    lines.append("  • Related party transactions outside ordinary course")

    if ts.token_rights.enabled:
        lines.append("")
        lines.append("TOKEN RIGHTS")
        lines.append("-" * 40)
        lines.append(
            f"Token Floor:            {ts.token_rights.token_floor_percent}% of Launch Supply"
        )
        lines.append("Token Allocation:       Pro rata share (fully-diluted) of Insider allocation")
        lines.append("Lockup:                 No more restrictive than Company/Insider lockup")
        if ts.token_rights.pro_rata_on_tokens:
            lines.append("Pro Rata on Tokens:     Yes")
        if ts.token_rights.side_letter:
            lines.append("Side Letter:            Yes")
        if ts.token_rights.warrant:
            lines.append("Token Warrant:          Yes")

    lines.append("")
    lines.append("OTHER TERMS")
    lines.append("-" * 40)
    lines.append(f"Legal Fee Cap:          {_format_money(ts.legal_fee_cap)}")
    lines.append(f"Exclusivity:            {ts.exclusivity_days} days")
    lines.append(f"Governing Law:          {ts.governing_law}")
    lines.append("Documentation:          Based on 2025 NVCA forms")

    if ts.custom_terms:
        lines.append("")
        lines.append("ADDITIONAL TERMS")
        lines.append("-" * 40)
        lines.append(ts.custom_terms)

    lines.append("")
    lines.append("=" * 60)
    lines.append("")
    lines.append("This term sheet is non-binding except for the No-Shop and")
    lines.append("Confidentiality provisions.")
    lines.append("")
    lines.append("Acknowledged and agreed:")
    lines.append("")
    lines.append("PARADIGM                         " + ts.company_name.upper())
    lines.append("")
    lines.append("By: _________________________    By: _________________________")
    lines.append("Name:                            Name:")
    lines.append("Title:                           Title:")
    lines.append("Date:                            Date:")

    return "\n".join(lines)


def generate_term_sheet_docx(ts: TermSheet) -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(ts.company_name.upper())
    run.bold = True
    run.font.size = Pt(14)

    if ts.instrument_type == InstrumentType.SAFE:
        subhead = "SIMPLE AGREEMENT FOR FUTURE EQUITY (SAFE)"
    elif ts.instrument_type == InstrumentType.PRICED:
        series = ts.series or "A"
        subhead = f"SERIES {series.upper()} PREFERRED STOCK FINANCING"
    else:
        subhead = "CONVERTIBLE PROMISSORY NOTE"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subhead)
    run.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("SUMMARY OF PROPOSED TERMS")

    doc.add_paragraph()

    def add_section(title: str):
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.underline = True

    def add_term(label: str, value: str):
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)

    add_section("Investment & Post-Money Valuation")
    investment_text = f'Paradigm Fund LP (together with its affiliates, "Paradigm") to invest {_format_money(ts.investment_amount)}'

    if ts.instrument_type == InstrumentType.PRICED:
        if ts.post_money_valuation:
            investment_text += (
                f" at a {_format_money(ts.post_money_valuation)} post-money valuation"
            )
            ownership = _calculate_ownership(
                ts.investment_amount, post_money=ts.post_money_valuation
            )
            investment_text += f" (including conversion of all convertible securities and an unallocated option pool equal to {ts.option_pool_percent}% of the post-money fully diluted capitalization)"
            investment_text += f', such that post-closing Paradigm will own {ownership:.1f}% of the fully diluted capitalization of {ts.company_name} (the "Company").'
    elif ts.instrument_type == InstrumentType.SAFE:
        if ts.valuation_cap:
            investment_text += f" via a SAFE with a {_format_money(ts.valuation_cap)} valuation cap"
        if ts.discount_percent:
            investment_text += f" and {ts.discount_percent}% discount"
        investment_text += "."
    elif ts.instrument_type == InstrumentType.CONVERTIBLE_NOTE:
        if ts.valuation_cap:
            investment_text += (
                f" via convertible note with a {_format_money(ts.valuation_cap)} valuation cap"
            )
        if ts.discount_percent:
            investment_text += f" and {ts.discount_percent}% discount"
        investment_text += "."

    doc.add_paragraph(investment_text)

    if ts.instrument_type == InstrumentType.PRICED:
        add_section("Securities")
        series = ts.series or "A"
        securities_text = f'Series {series} Preferred Stock (together with other series of Preferred Stock, the "Preferred Stock") with standard non-cumulative dividends in preference of Common Stock, {ts.liquidation_preference} liquidation preference and {ts.anti_dilution} antidilution protection (subject to limited exclusions), that is convertible to Common Stock upon the earlier of (i) the election of Preferred Majority (as defined below) or (ii) the consummation of an underwritten public offering with net proceeds greater than $100M.'
        doc.add_paragraph(securities_text)

    add_section("Board and Voting Rights")
    if ts.board_rights == BoardRights.SEAT:
        board_text = (
            "One director to be elected by the Series Preferred Stock and designated by Paradigm. "
        )
    elif ts.board_rights == BoardRights.SEAT_AND_OBSERVER:
        board_text = "One director to be elected by the Series Preferred Stock and designated by Paradigm. Company shall also invite a representative of Paradigm to attend all meetings of the Board in a nonvoting observer capacity and provide such representative copies of all notices, minutes, consents, and other materials provided to its directors. "
    elif ts.board_rights == BoardRights.OBSERVER:
        board_text = "Company shall invite a representative of Paradigm to attend all meetings of the Board in a nonvoting observer capacity and provide such representative copies of all notices, minutes, consents, and other materials provided to its directors. "
    else:
        board_text = ""
    board_text += 'Preferred Stock voting thresholds to be set such that Paradigm\'s consent is required (the "Preferred Majority").'
    doc.add_paragraph(board_text)

    add_section("Protective Provisions")
    doc.add_paragraph(
        "Consent of the Preferred Majority required for standard NVCA protective provisions and certain additional protective provisions, including:"
    )
    provisions = [
        "incurrence of indebtedness or issuance of debt securities greater than $1M",
        "creation of any new equity compensation plan or increase the number of shares available for issuance pursuant to such plans",
        "any sale, assignment, license, pledge or encumbrance of material technology or intellectual property of the Company",
        'the creation, reservation, sale, distribution, issuance or other disposition of any tokens ("Tokens")',
        "any interested or related party transactions other than transactions entered into in the ordinary course of business on an arms-length basis",
    ]
    for i, provision in enumerate(provisions, 1):
        doc.add_paragraph(f"({i}) {provision};")

    add_section("Other Rights")
    if ts.pro_rata_rights:
        rights_text = "Customary NVCA investor rights, including information rights and pro rata rights (including overallotment) for Major Investors (which shall only include Paradigm), registration rights for all Investors, drag along provision, and all 1% Common stockholder's (including founder(s)) equity and Tokens will be subject to ROFR and co-sale rights."
    else:
        rights_text = "Customary NVCA investor rights, including information rights, registration rights for all Investors, and drag along provision."
    doc.add_paragraph(rights_text)

    if ts.token_rights.enabled:
        add_section("Token Rights")
        token_text = f'For any Tokens (other than non-fungible tokens or other similar assets developed in the ordinary course of business) useable or accessible in or through a blockchain-based game or application created by the Company, founder or affiliates, Paradigm will receive its pro rata share (on a fully-diluted basis, as of network launch) of the total number of Tokens (the "Launch Supply") allocated to or reserved for the Company, the Company\'s officers, directors, employees, consultants, stockholders and any convertible instrument holders (collectively, the "Insiders"). The Launch Supply shall be at least {ts.token_rights.token_floor_percent}% of the total number of Tokens issuable for such network.'
        doc.add_paragraph(token_text)
        doc.add_paragraph(
            "If an inflationary event (the creation of additional Tokens following network launch) occurs, Paradigm will receive its pro rata share (on a fully-diluted basis, as of the date of such inflationary event) of the total number of inflationary tokens allocated to or reserved for Insiders in connection with such inflationary event."
        )
        doc.add_paragraph(
            "Any lockup schedule on such Tokens shall be no more restrictive than the schedule applicable to Tokens issued to the Company or Insiders."
        )

    add_section("Vesting")
    doc.add_paragraph(
        "Founder vesting subject to due diligence. Standard 4-year monthly vesting with one year cliff for all employees, beginning on first day of employment."
    )

    add_section("Documentation; Legal Fees")
    doc.add_paragraph(
        f"Company counsel to draft documentation based on 2025 NVCA forms. Company will pay the reasonable legal fees incurred by Paradigm's counsel up to {_format_money(ts.legal_fee_cap)}."
    )

    add_section("No-Shop; Confidentiality")
    doc.add_paragraph(
        f"The Company and the founders agree that they will not, for a period of {ts.exclusivity_days} days from the date these terms are accepted, take any action to solicit, initiate, encourage or assist the submission of any proposal, negotiation or offer from any person or entity other than Paradigm relating to the sale or issuance of any of the capital stock of the Company."
    )
    doc.add_paragraph(
        "The Company will not disclose the terms of this Term Sheet to any person other than officers, members of the Board and the Company's accountants and attorneys and other potential investors acceptable to Paradigm, without the written consent of Paradigm."
    )

    if ts.custom_terms:
        add_section("Additional Terms")
        doc.add_paragraph(ts.custom_terms)

    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer.add_run(
        "Except for the No-Shop; Confidentiality provision set forth above, this term sheet is non-binding and is intended solely to be a summary of the terms that are currently proposed by the parties."
    ).italic = True

    doc.add_paragraph()
    doc.add_paragraph(
        "Please indicate your acceptance of this term sheet by signing below and returning an executed copy."
    )

    doc.add_paragraph()
    doc.add_paragraph("Acknowledged and agreed:")

    doc.add_paragraph()
    table = doc.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "PARADIGM"
    table.cell(0, 1).text = ts.company_name.upper()
    table.cell(1, 0).text = "By: _________________________"
    table.cell(1, 1).text = "By: _________________________"
    table.cell(2, 0).text = "Name:"
    table.cell(2, 1).text = "Name:"
    table.cell(3, 0).text = "Title:                    Date:"
    table.cell(3, 1).text = "Title:                    Date:"

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def generate_draft_email(ts: TermSheet, dri_name: str = "") -> str:
    if ts.instrument_type == InstrumentType.SAFE:
        valuation_line = (
            f"Valuation Cap: {_format_money(ts.valuation_cap)}" if ts.valuation_cap else ""
        )
        instrument = "SAFE"
    elif ts.instrument_type == InstrumentType.PRICED:
        if ts.post_money_valuation:
            valuation_line = f"Post-Money Valuation: {_format_money(ts.post_money_valuation)}"
        elif ts.pre_money_valuation:
            valuation_line = f"Pre-Money Valuation: {_format_money(ts.pre_money_valuation)}"
        else:
            valuation_line = ""
        series = ts.series or "A"
        instrument = f"Series {series} Preferred Stock"
    else:
        valuation_line = (
            f"Valuation Cap: {_format_money(ts.valuation_cap)}" if ts.valuation_cap else ""
        )
        instrument = "Convertible Note"

    founder_name = ts.founder_name or "[Founder]"
    dri = dri_name or ts.dri_name or "[DRI Name]"

    email = f"""Subject: {ts.company_name} - Term Sheet

Hi {founder_name},

Following our conversation, please find attached our proposed term sheet for {ts.company_name}.

Key terms:
- Investment: {_format_money(ts.investment_amount)}
- {valuation_line}
- Instrument: {instrument}"""

    if ts.token_rights.enabled:
        email += f"\n- Token Rights: Yes ({ts.token_rights.token_floor_percent}% floor)"

    email += f"""

Please review and let us know if you have any questions. We're excited about the opportunity to partner with you.

Best,
{dri}
Paradigm"""

    return email
