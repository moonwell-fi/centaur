"""Template-based term sheet generation with strict format compliance.

This module ensures all term sheets follow the exact Paradigm template format.
Key rules enforced:
1. Company name in ALL CAPS in title
2. NO BRACKETS - all placeholders must be filled
3. $1M debt threshold in protective provisions
4. 2025 NVCA forms reference
5. $75K legal fee cap default
6. Exact section ordering from Paradigm template
"""

from io import BytesIO

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .models import TermSheet, InstrumentType, BoardRights


# Strict template section order - must match Paradigm template exactly
TEMPLATE_SECTIONS = [
    "Investment & Post-Money Valuation",
    "Securities",
    "Board and Voting Rights",
    "Protective Provisions",
    "Other Rights",
    "Token Rights",  # Only included if token_rights.enabled
    "Vesting",
    "Documentation; Legal Fees",
    "No-Shop; Confidentiality",
]


def _format_money(amount: float) -> str:
    """Format money without brackets or placeholders."""
    if amount >= 1_000_000_000:
        billions = amount / 1_000_000_000
        if billions == int(billions):
            return f"${int(billions)}B"
        return f"${billions:.1f}B"
    elif amount >= 1_000_000:
        millions = amount / 1_000_000
        if millions == int(millions):
            return f"${int(millions)}M"
        return f"${millions:.1f}M"
    elif amount >= 1_000:
        return f"${amount / 1_000:.0f}K"
    else:
        return f"${amount:,.0f}"


def _format_money_full(amount: float) -> str:
    """Format money as full dollar amount (e.g., $75,000)."""
    return f"${amount:,.0f}"


def _calculate_ownership(
    investment: float,
    post_money: float | None = None,
    pre_money: float | None = None,
) -> float:
    """Calculate ownership percentage."""
    if post_money:
        return (investment / post_money) * 100
    elif pre_money:
        return (investment / (pre_money + investment)) * 100
    return 0.0


def _validate_no_brackets(text: str) -> str:
    """Ensure no brackets remain in output text."""
    if "[" in text or "]" in text:
        raise ValueError(
            f"Template output contains brackets - all placeholders must be filled: {text[:100]}"
        )
    return text


class TemplateGenerator:
    """Generates term sheets using strict template format."""

    def __init__(self, ts: TermSheet):
        self.ts = ts
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Set up document styles matching Paradigm template."""
        style = self.doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)

    def _add_title(self):
        """Add document title - company name MUST be in ALL CAPS."""
        # Title line 1: Company name in ALL CAPS
        heading = self.doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company_upper = self.ts.company_name.upper()
        run = heading.add_run(company_upper)
        run.bold = True
        run.font.size = Pt(14)

        # Title line 2: Instrument type
        if self.ts.instrument_type == InstrumentType.SAFE:
            subhead = "SIMPLE AGREEMENT FOR FUTURE EQUITY (SAFE)"
        elif self.ts.instrument_type == InstrumentType.PRICED:
            series = self.ts.series or "A"
            subhead = f"SERIES {series.upper()} PREFERRED STOCK FINANCING"
        else:
            subhead = "CONVERTIBLE PROMISSORY NOTE"

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subhead)
        run.bold = True
        run.font.size = Pt(12)

        # Subtitle
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("SUMMARY OF PROPOSED TERMS")

        self.doc.add_paragraph()

    def _add_section_header(self, title: str):
        """Add a section header with underline."""
        p = self.doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.underline = True

    def _add_paragraph(self, text: str):
        """Add a paragraph, validating no brackets remain."""
        validated = _validate_no_brackets(text)
        self.doc.add_paragraph(validated)

    def _add_numbered_item(self, number: int, text: str):
        """Add a numbered item in protective provisions style."""
        validated = _validate_no_brackets(text)
        self.doc.add_paragraph(f"({number}) {validated};")

    def _section_investment(self):
        """Investment & Post-Money Valuation section."""
        self._add_section_header("Investment & Post-Money Valuation")

        ts = self.ts
        company = ts.company_name

        if ts.instrument_type == InstrumentType.PRICED:
            if ts.post_money_valuation:
                ownership = _calculate_ownership(
                    ts.investment_amount, post_money=ts.post_money_valuation
                )
                text = (
                    f'Paradigm Fund LP (together with its affiliates, "Paradigm") to invest '
                    f"{_format_money(ts.investment_amount)} at a {_format_money(ts.post_money_valuation)} "
                    f"post-money valuation (including conversion of all convertible securities and an "
                    f"unallocated option pool equal to {ts.option_pool_percent:.0f}% of the post-money "
                    f"fully diluted capitalization), such that post-closing Paradigm will own "
                    f'{ownership:.1f}% of the fully diluted capitalization of {company} (the "Company").'
                )
            elif ts.pre_money_valuation:
                ownership = _calculate_ownership(
                    ts.investment_amount, pre_money=ts.pre_money_valuation
                )
                text = (
                    f'Paradigm Fund LP (together with its affiliates, "Paradigm") to invest '
                    f"{_format_money(ts.investment_amount)} at a {_format_money(ts.pre_money_valuation)} "
                    f"pre-money valuation (including conversion of all convertible securities and an "
                    f"unallocated option pool equal to {ts.option_pool_percent:.0f}% of the post-money "
                    f"fully diluted capitalization), such that post-closing Paradigm will own "
                    f'{ownership:.1f}% of the fully diluted capitalization of {company} (the "Company").'
                )
            else:
                text = (
                    f'Paradigm Fund LP (together with its affiliates, "Paradigm") to invest '
                    f'{_format_money(ts.investment_amount)} in {company} (the "Company").'
                )
        elif ts.instrument_type == InstrumentType.SAFE:
            text = f'Paradigm Fund LP (together with its affiliates, "Paradigm") to invest {_format_money(ts.investment_amount)}'
            if ts.valuation_cap:
                text += f" via a SAFE with a {_format_money(ts.valuation_cap)} valuation cap"
            if ts.discount_percent:
                text += f" and {ts.discount_percent:.0f}% discount"
            text += f' in {company} (the "Company").'
        else:  # Convertible Note
            text = f'Paradigm Fund LP (together with its affiliates, "Paradigm") to invest {_format_money(ts.investment_amount)}'
            if ts.valuation_cap:
                text += (
                    f" via convertible note with a {_format_money(ts.valuation_cap)} valuation cap"
                )
            if ts.discount_percent:
                text += f" and {ts.discount_percent:.0f}% discount"
            text += f' in {company} (the "Company").'

        self._add_paragraph(text)

    def _section_securities(self):
        """Securities section - only for priced rounds."""
        if self.ts.instrument_type != InstrumentType.PRICED:
            return

        self._add_section_header("Securities")

        series = self.ts.series or "A"
        text = (
            f"Series {series} Preferred Stock (together with other series of Preferred Stock, "
            f'the "Preferred Stock") with standard non-cumulative dividends in preference of '
            f"Common Stock, {self.ts.liquidation_preference} liquidation preference and "
            f"{self.ts.anti_dilution} antidilution protection (subject to limited exclusions), "
            f"that is convertible to Common Stock upon the earlier of (i) the election of "
            f"Preferred Majority (as defined below) or (ii) the consummation of an underwritten "
            f"public offering with net proceeds greater than $100M."
        )
        self._add_paragraph(text)

    def _section_board(self):
        """Board and Voting Rights section."""
        self._add_section_header("Board and Voting Rights")

        ts = self.ts
        if ts.board_rights == BoardRights.SEAT:
            board_text = "One director to be elected by the Series Preferred Stock and designated by Paradigm. "
        elif ts.board_rights == BoardRights.SEAT_AND_OBSERVER:
            board_text = (
                "One director to be elected by the Series Preferred Stock and designated by Paradigm. "
                "Company shall also invite a representative of Paradigm to attend all meetings of the "
                "Board in a nonvoting observer capacity and provide such representative copies of all "
                "notices, minutes, consents, and other materials provided to its directors. "
            )
        elif ts.board_rights == BoardRights.OBSERVER:
            board_text = (
                "Company shall invite a representative of Paradigm to attend all meetings of the Board "
                "in a nonvoting observer capacity and provide such representative copies of all notices, "
                "minutes, consents, and other materials provided to its directors. "
            )
        else:
            board_text = ""

        board_text += (
            "Preferred Stock voting thresholds to be set such that Paradigm's consent is "
            'required (the "Preferred Majority").'
        )

        self._add_paragraph(board_text)

    def _section_protective_provisions(self):
        """Protective Provisions section - MUST use $1M debt threshold, NO brackets."""
        self._add_section_header("Protective Provisions")

        self._add_paragraph(
            "Consent of the Preferred Majority required for standard NVCA protective provisions "
            "and certain additional protective provisions, including:"
        )

        # Standard protective provisions - debt threshold is ALWAYS $1M, no brackets
        provisions = [
            "incurrence of indebtedness or issuance of debt securities greater than $1,000,000",
            "creation of any new equity compensation plan or increase the number of shares available for issuance pursuant to such plans",
            "any sale, assignment, license, pledge or encumbrance of material technology or intellectual property of the Company",
            'the creation, reservation, sale, distribution, issuance or other disposition of any tokens ("Tokens")',
            "any interested or related party transactions other than transactions entered into in the ordinary course of business on an arms-length basis",
        ]

        for i, provision in enumerate(provisions, 1):
            self._add_numbered_item(i, provision)

    def _section_other_rights(self):
        """Other Rights section."""
        self._add_section_header("Other Rights")

        if self.ts.pro_rata_rights:
            text = (
                "Customary NVCA investor rights, including information rights and pro rata rights "
                "(including overallotment) for Major Investors (which shall only include Paradigm), "
                "registration rights for all Investors, drag along provision, and all 1% Common "
                "stockholder's (including founder(s)) equity and Tokens will be subject to ROFR "
                "and co-sale rights."
            )
        else:
            text = (
                "Customary NVCA investor rights, including information rights, registration rights "
                "for all Investors, and drag along provision."
            )

        self._add_paragraph(text)

    def _section_token_rights(self):
        """Token Rights section - only included if enabled."""
        if not self.ts.token_rights.enabled:
            return

        self._add_section_header("Token Rights")

        floor_pct = self.ts.token_rights.token_floor_percent
        text = (
            f"For any Tokens (other than non-fungible tokens or other similar assets developed "
            f"in the ordinary course of business) useable or accessible in or through a blockchain-based "
            f"game or application created by the Company, founder or affiliates, Paradigm will receive "
            f"its pro rata share (on a fully-diluted basis, as of network launch) of the total number "
            f'of Tokens (the "Launch Supply") allocated to or reserved for the Company, the Company\'s '
            f"officers, directors, employees, consultants, stockholders and any convertible instrument "
            f'holders (collectively, the "Insiders"). The Launch Supply shall be at least '
            f"{floor_pct:.0f}% of the total number of Tokens issuable for such network."
        )
        self._add_paragraph(text)

        self._add_paragraph(
            "If an inflationary event (the creation of additional Tokens following network launch) "
            "occurs, Paradigm will receive its pro rata share (on a fully-diluted basis, as of the "
            "date of such inflationary event) of the total number of inflationary tokens allocated "
            "to or reserved for Insiders in connection with such inflationary event."
        )

        self._add_paragraph(
            "Any lockup schedule on such Tokens shall be no more restrictive than the schedule "
            "applicable to Tokens issued to the Company or Insiders."
        )

    def _section_vesting(self):
        """Vesting section."""
        self._add_section_header("Vesting")

        self._add_paragraph(
            "Founder vesting subject to due diligence. Standard 4-year monthly vesting with "
            "one year cliff for all employees, beginning on first day of employment."
        )

    def _section_documentation(self):
        """Documentation; Legal Fees section - MUST reference 2025 NVCA forms."""
        self._add_section_header("Documentation; Legal Fees")

        fee_cap = _format_money_full(self.ts.legal_fee_cap)
        self._add_paragraph(
            f"Company counsel to draft documentation based on 2025 NVCA forms. Company will pay "
            f"the reasonable legal fees incurred by Paradigm's counsel up to {fee_cap}."
        )

    def _section_no_shop(self):
        """No-Shop; Confidentiality section."""
        self._add_section_header("No-Shop; Confidentiality")

        days = self.ts.exclusivity_days
        self._add_paragraph(
            f"The Company and the founders agree that they will not, for a period of {days} days "
            f"from the date these terms are accepted, take any action to solicit, initiate, encourage "
            f"or assist the submission of any proposal, negotiation or offer from any person or entity "
            f"other than Paradigm relating to the sale or issuance of any of the capital stock of the Company."
        )

        self._add_paragraph(
            "The Company will not disclose the terms of this Term Sheet to any person other than "
            "officers, members of the Board and the Company's accountants and attorneys and other "
            "potential investors acceptable to Paradigm, without the written consent of Paradigm."
        )

    def _add_custom_terms(self):
        """Add custom terms if present."""
        if self.ts.custom_terms:
            self._add_section_header("Additional Terms")
            self._add_paragraph(self.ts.custom_terms)

    def _add_disclaimer(self):
        """Add non-binding disclaimer."""
        self.doc.add_paragraph()
        disclaimer = self.doc.add_paragraph()
        run = disclaimer.add_run(
            "Except for the No-Shop; Confidentiality provision set forth above, this term sheet "
            "is non-binding and is intended solely to be a summary of the terms that are currently "
            "proposed by the parties."
        )
        run.italic = True

    def _add_signature_block(self):
        """Add signature block - NO bullet points, NO DRI checklist."""
        self.doc.add_paragraph()
        self._add_paragraph(
            "Please indicate your acceptance of this term sheet by signing below and returning "
            "an executed copy."
        )

        self.doc.add_paragraph()
        self._add_paragraph("Acknowledged and agreed:")

        self.doc.add_paragraph()
        table = self.doc.add_table(rows=4, cols=2)
        table.cell(0, 0).text = "PARADIGM"
        table.cell(0, 1).text = self.ts.company_name.upper()
        table.cell(1, 0).text = "By: _________________________"
        table.cell(1, 1).text = "By: _________________________"
        table.cell(2, 0).text = "Name:"
        table.cell(2, 1).text = "Name:"
        table.cell(3, 0).text = "Title:                    Date:"
        table.cell(3, 1).text = "Title:                    Date:"

    def generate(self) -> bytes:
        """Generate the complete term sheet document.

        Returns:
            bytes: The generated .docx file as bytes
        """
        # Build document in strict template order
        self._add_title()
        self._section_investment()
        self._section_securities()
        self._section_board()
        self._section_protective_provisions()
        self._section_other_rights()
        self._section_token_rights()
        self._section_vesting()
        self._section_documentation()
        self._section_no_shop()
        self._add_custom_terms()
        self._add_disclaimer()
        self._add_signature_block()

        # Save to bytes
        buffer = BytesIO()
        self.doc.save(buffer)
        return buffer.getvalue()


def generate_term_sheet_docx_strict(ts: TermSheet) -> bytes:
    """Generate a term sheet with strict template compliance.

    This is the preferred method for generating term sheets as it:
    1. Enforces exact Paradigm template format
    2. Validates no brackets remain in output
    3. Uses correct defaults ($75K legal, 2025 NVCA, $1M debt)
    4. Follows exact section ordering

    Args:
        ts: The TermSheet data model

    Returns:
        bytes: The generated .docx file
    """
    generator = TemplateGenerator(ts)
    return generator.generate()
