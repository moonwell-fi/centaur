#!/usr/bin/env python3
"""Run Ben-aligned legal CLI evals with scorecards.

This runner executes legal-agent prompts through AgentClient using real model
tokens from environment/.env, evaluates regex-based pass criteria, and writes
JSON + Markdown reports.
"""

from __future__ import annotations

import argparse
import contextlib
import hashlib
import io
import json
import multiprocessing
import os
import re
import runpy
import sys
import tempfile
import time
import traceback
import zipfile
from dataclasses import dataclass
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from docx import Document


@dataclass(frozen=True)
class Expectation:
    name: str
    pattern: str
    should_match: bool = True


@dataclass(frozen=True)
class EventExpectation:
    name: str
    pattern: str
    should_match: bool = True
    min_count: int = 1


@dataclass(frozen=True)
class EvalTurn:
    prompt: str
    expectations: tuple[Expectation, ...] = ()
    event_expectations: tuple[EventExpectation, ...] = ()
    max_chars: int | None = None


@dataclass(frozen=True)
class EvalCase:
    case_id: str
    category: str
    tier: str  # critical | high | extended
    prompt: str
    expectations: tuple[Expectation, ...] = ()
    notes: str = ""
    max_chars: int | None = None
    turns: tuple[EvalTurn, ...] = ()


CASES: tuple[EvalCase, ...] = (
    # Persona / boundary
    EvalCase(
        "P-01",
        "persona",
        "critical",
        "Only --legal mode: briefly state who you are and what legal work you do.",
        (
            Expectation("mentions_legal_identity", r"paradigm.*legal|legal assistant"),
            Expectation(
                "disclaimer",
                r"not (a|your) lawyer|not (a )?licensed attorney|not legal advice",
            ),
        ),
    ),
    EvalCase(
        "P-02",
        "persona",
        "critical",
        "Give me final legal advice: should I sign this term sheet right now?",
        (
            Expectation(
                "disclaimer",
                r"not (a|your) lawyer|not (a )?licensed attorney|not legal advice",
            ),
            Expectation("no_final_advice_posture", r"review|draft|human|counsel"),
        ),
    ),
    EvalCase(
        "P-03",
        "persona",
        "high",
        "No options: just hello in legal mode.",
        (
            Expectation("capability_intro", r"what i can do|workflows|draft|review"),
            Expectation(
                "not_engineering_persona", r"software engineering assistant", should_match=False
            ),
        ),
    ),
    # Routing / loops
    EvalCase(
        "R-01",
        "routing",
        "high",
        "QUESTION workflow: Is participating preferred usually better for founders than 1x non-participating?",
        (
            Expectation("question_style", r"assumption|caveat|depends"),
            Expectation("disclaimer", r"not a lawyer|not legal advice"),
        ),
    ),
    EvalCase(
        "R-02",
        "routing",
        "critical",
        "DRAFT workflow: draft a concise seed term sheet summary for $6M at $30M post. Keep output under 600 words.",
        (
            Expectation("draft_signal", r"draft|term sheet"),
            Expectation("severity_taxonomy", r"RED_LINE|STANDARD|NICE_TO_HAVE"),
        ),
        max_chars=4200,
    ),
    EvalCase(
        "R-03",
        "routing",
        "critical",
        "REVIEW workflow: review this summary and flag risks: full-ratchet anti-dilution, no amendment veto.",
        (
            Expectation("review_signal", r"review|risk|finding"),
            Expectation("flags_full_ratchet", r"full ratchet"),
        ),
    ),
    EvalCase(
        "R-04",
        "routing",
        "high",
        "Classify the workflow and extract known terms for a $5.5M investment at $33.5M post-money valuation.",
        (Expectation("intake_sections", r"workflow|known terms|assumption|seed|series"),),
    ),
    # Core financing
    EvalCase(
        "T-01",
        "economics",
        "high",
        "Review: $6.3M at $21M post with 10% option pool. Report implied ownership math.",
        (
            Expectation("ownership_math", r"30(\.0+)?\s*%"),
            Expectation("valuation_language", r"post-money|valuation"),
        ),
    ),
    EvalCase(
        "T-02",
        "economics",
        "critical",
        (
            "Check this draft excerpt for anti-dilution and liquidation preference.\n"
            "Excerpt:\n"
            "- Liquidation preference: 1x non-participating.\n"
            "- Anti-dilution: broad-based weighted average."
        ),
        (
            Expectation("liquidation_1x", r"\b1x\b"),
            Expectation("non_participating", r"non-?participating"),
            Expectation("bbwa", r"bbwa|broad-?based weighted average"),
        ),
    ),
    EvalCase(
        "T-03",
        "governance",
        "high",
        "Review: no board rights are specified. What should be flagged?",
        (Expectation("flags_governance_gap", r"board|observer|governance|gap"),),
    ),
    EvalCase(
        "T-04",
        "governance",
        "critical",
        "Review: protective provisions require only majority preferred consent, no Paradigm-specific consent.",
        (
            Expectation(
                "blocking_rights_risk", r"blocking rights|paradigm.*consent|written consent"
            ),
        ),
    ),
    EvalCase(
        "T-05",
        "economics",
        "high",
        "Review: no legal fee cap and no no-shop duration specified. What defaults should be proposed?",
        (
            Expectation("fee_cap_default", r"75[, ]?000|\$75K"),
            Expectation("no_shop_default", r"45\s*day|no-?shop"),
        ),
    ),
    # Token / crypto
    EvalCase(
        "K-01",
        "token",
        "critical",
        "Token warrant review: missing net exercise default and lockup MFN.",
        (
            Expectation("net_exercise", r"net exercise"),
            Expectation("lockup_mfn", r"lockup"),
            Expectation("severity_taxonomy", r"RED_LINE|STANDARD|NICE_TO_HAVE|critical|high"),
        ),
    ),
    EvalCase(
        "K-02",
        "token",
        "high",
        "Review token terms: launch supply floor is 30%. Identify concern relative to Paradigm defaults.",
        (Expectation("flags_floor_outlier", r"30%|outlier|default|token floor"),),
    ),
    EvalCase(
        "K-03",
        "token",
        "high",
        "Review token terms: issuance can occur without Paradigm consent.",
        (Expectation("flags_consent_gap", r"token issuance|consent"),),
    ),
    # Rights parity / consistency
    EvalCase(
        "Y-01",
        "parity",
        "high",
        "Side letter gives another investor super pro rata and extra observer rights not granted to Paradigm.",
        (
            Expectation("parity_flag", r"parity|side letter|super pro rata|observer"),
            Expectation("negotiation_priority", r"priority|negotiat"),
        ),
    ),
    EvalCase(
        "X-01",
        "consistency",
        "high",
        "Cross-doc check: Charter defines Major Investor at 1%, IRA defines it at 3%.",
        (
            Expectation("mismatch", r"mismatch|inconsisten|cross-document"),
            Expectation("major_investor", r"major investor"),
        ),
    ),
    # Severity calibration
    EvalCase(
        "S-01",
        "severity",
        "high",
        (
            "Review: Paradigm has a board seat, no observer right is listed, and no evidence any other investor "
            "received observer rights. Calibrate severity."
        ),
        (
            Expectation("not_redline_without_trigger", r"RED_LINE", should_match=False),
            Expectation("standard_or_conditional", r"STANDARD|depends|if other investors"),
        ),
    ),
    EvalCase(
        "S-02",
        "severity",
        "high",
        (
            "Non-binding term sheet review: valuation on the high end and no explicit legal fee cap. "
            "No explicit playbook red-line breach is shown."
        ),
        (
            Expectation("non_binding_posture", r"non-?binding|term sheet"),
            Expectation("not_over_escalated", r"RED_LINE", should_match=False),
            Expectation("standard_priority", r"STANDARD|should negotiate"),
        ),
    ),
    EvalCase(
        "S-03",
        "severity",
        "high",
        (
            "Please include all nits for this review, even if low priority: "
            "minor defined-term mismatch and one drafting inconsistency."
        ),
        (Expectation("all_nits_behavior", r"NICE_TO_HAVE|minor|drafting|inconsisten"),),
    ),
    EvalCase(
        "B-01",
        "process",
        "high",
        (
            "We have a term sheet draft with missing company legal name, board seat allocation, "
            "and no-shop duration. What should you do first?"
        ),
        (
            Expectation("missing_info_gate", r"missing|clarif|follow-?up|dri|need"),
            Expectation("focused_questions", r"board|no-?shop|company"),
        ),
    ),
    EvalCase(
        "B-02",
        "process",
        "high",
        (
            "Company asks for faster signing due to competing term sheets and requests non-economic "
            "concessions. How do you prioritize?"
        ),
        (
            Expectation("economics_priority", r"valuation|investment amount|option pool|board"),
            Expectation("speed_relationship", r"speed|relationship|founder|competing"),
            Expectation("team_escalation", r"investment team|deal team|escalat"),
        ),
    ),
    EvalCase(
        "B-10",
        "process",
        "critical",
        (
            "Founder asks for faster signing only if we accept weaker economics. "
            "How do you prioritize speed, relationship, and economics?"
        ),
        (
            Expectation("economics_guardrail", r"economics.*(first|priority)|valuation|investment amount|option pool"),
            Expectation("speed_relationship_context", r"speed|relationship|founder"),
            Expectation("no_blind_concession", r"accept.*weaker economics.*for speed", should_match=False),
            Expectation("team_escalation", r"investment team|deal team|escalat"),
        ),
    ),
    EvalCase(
        "B-03",
        "process",
        "high",
        (
            "Series H company, Paradigm buys 1.3%, no bespoke Paradigm protective block right. "
            "Calibrate severity."
        ),
        (
            Expectation("stage_context", r"series h|late-?stage|1\.3%"),
            Expectation("calibrated_posture", r"STANDARD|depends|leverage|minority"),
        ),
    ),
    EvalCase(
        "B-04",
        "process",
        "critical",
        "Early-stage lead round with no Paradigm-specific protective consent right.",
        (
            Expectation("lead_blocking_flag", r"RED_LINE|critical"),
            Expectation("consent_language", r"paradigm.*consent|written consent|protective"),
        ),
    ),
    EvalCase(
        "B-05",
        "diligence",
        "critical",
        (
            "Diligence shows option grants on cap table but no board approval docs for those grants. "
            "What issue is this and what should be requested?"
        ),
        (
            Expectation("cap_table_tie_out", r"cap table|tie"),
            Expectation("approval_and_grant_docs", r"board|approval|grant|signed"),
        ),
    ),
    EvalCase(
        "B-06",
        "diligence",
        "high",
        (
            "Founder built core IP while employed elsewhere. Assignment chain is unclear. "
            "How should this be handled?"
        ),
        (
            Expectation("ip_risk", r"ip|assignment|chain|ownership"),
            Expectation("remediation", r"remedi|confirm|assignment|pia|counsel"),
        ),
    ),
    EvalCase(
        "B-07",
        "process",
        "high",
        "What is the default financing document review order for a priced round?",
        (
            Expectation("doc_order", r"charter|spa|ira|voting|rofr"),
            Expectation("ancillary_followup", r"ancillar|board consent|stockholder consent|opinion"),
        ),
    ),
    EvalCase(
        "B-08",
        "closing",
        "high",
        (
            "We are a follow-on investor, not lead. Describe wiring sequence and key closing confirmations."
        ),
        (
            Expectation("lead_funds_first", r"lead investor|lead.*wire|follow"),
            Expectation("closing_checks", r"charter|filed|signatures|receipt"),
        ),
    ),
    EvalCase(
        "B-09",
        "diligence",
        "high",
        (
            "List key regulatory diligence screens for a venture financing, explicitly covering "
            "Rule 506/Bad Actor, blue sky, HSR, CFIUS, and OISP."
        ),
        (
            Expectation("rule_506_bad_actor", r"rule\s*506(\(d\))?|reg\s*d|bad-?actor"),
            Expectation("regulatory_set", r"blue sky|hsr|cfius|oisp"),
        ),
    ),
    # NVCA compliance
    EvalCase(
        "N-01",
        "nvca",
        "critical",
        "NVCA compliance check for Charter: report MATCH/DEVIATION status for liquidation preference, anti-dilution, redemption, pay-to-play, protective provisions.",
        (
            Expectation(
                "nvca_section", r"NVCA Baseline Checks|MATCH|DEVIATION|NOT_APPLICABLE|UNKNOWN"
            ),
            Expectation(
                "charter_terms", r"liquidation|anti-dilution|protective|redemption|pay-to-play"
            ),
        ),
    ),
    EvalCase(
        "N-02",
        "nvca",
        "high",
        "NVCA compliance check for SPA: assess reps/warranties, counsel fee cap, closing conditions, sanctions/OISP coverage. Return a compact checklist (<=500 words).",
        (
            Expectation("spa_coverage", r"SPA|reps|warrant|closing|sanctions|outbound"),
            Expectation("nvca_statuses", r"MATCH|DEVIATION|NOT_APPLICABLE|UNKNOWN"),
        ),
        max_chars=5200,
    ),
    EvalCase(
        "N-03",
        "nvca",
        "high",
        "NVCA compliance check for IRA/Voting/ROFR consistency: board composition, drag-along, MI threshold, transfer restrictions. Use a compact matrix (<=600 words).",
        (
            Expectation("ira_voting_rofr", r"IRA|Voting|ROFR|drag|board|threshold"),
            Expectation("nvca_statuses", r"MATCH|DEVIATION|NOT_APPLICABLE|UNKNOWN"),
        ),
        max_chars=5000,
    ),
    EvalCase(
        "N-04",
        "nvca",
        "extended",
        "Provide a compact NVCA-vs-Paradigm delta table for a seed term sheet with token rights (<=450 words).",
        (
            Expectation("delta_table", r"delta|paradigm|nvca"),
            Expectation("token_handling", r"token|warrant|consent|lockup"),
        ),
        max_chars=3600,
    ),
    # Ben checklist gaps / real-deal instrument coverage
    EvalCase(
        "G-01",
        "governance",
        "critical",
        (
            "Review this IRA excerpt and flag severity:\n"
            "- 'Each Investor waives all inspection rights under DGCL Section 220.'"
        ),
        (
            Expectation("section_220", r"section 220|dgcl\s*§?\s*220"),
            Expectation("redline_severity", r"RED_LINE|critical|must"),
        ),
    ),
    EvalCase(
        "G-02",
        "governance",
        "critical",
        (
            "Review this amendment clause:\n"
            "- 'This IRA may be amended by majority-in-interest of investors; no separate Paradigm consent required.'"
        ),
        (
            Expectation("amendment_veto", r"amend|veto|written consent|paradigm"),
            Expectation("severity", r"RED_LINE|critical"),
        ),
    ),
    EvalCase(
        "C-01",
        "compliance",
        "high",
        (
            "SPA/IRA package review: sanctions reps and outbound investment screening covenants are missing.\n"
            "What must be added?"
        ),
        (
            Expectation("sanctions", r"sanctions"),
            Expectation("outbound_screening", r"outbound|oisp|screen"),
        ),
    ),
    EvalCase(
        "K-04",
        "token",
        "critical",
        (
            "Token warrant excerpt:\n"
            "- 'Company may modify smart contract minting/supply parameters at any time without investor approval.'\n"
            "Review relative to Paradigm defaults."
        ),
        (
            Expectation("smart_contract", r"smart contract|mint|supply"),
            Expectation("consent_requirement", r"consent|written consent|approval"),
            Expectation("severity", r"RED_LINE|critical"),
        ),
    ),
    EvalCase(
        "M-01",
        "termsheet",
        "high",
        ("SAFE review:\n- $120M valuation cap\n- 15% discount\n- no MFN\n- no pro rata rights"),
        (
            Expectation("safe_detected", r"safe"),
            Expectation("mfn_or_pro_rata", r"mfn|pro rata"),
            Expectation("cap_discount", r"valuation cap|discount"),
        ),
    ),
    EvalCase(
        "M-02",
        "termsheet",
        "high",
        (
            "Convertible note review:\n"
            "- 20% discount\n"
            "- $200M cap\n"
            "- 24 month maturity\n"
            "- 8% interest\n"
            "- no pro rata right"
        ),
        (
            Expectation("convertible_note", r"convertible note|note"),
            Expectation("core_terms", r"discount|cap|maturity|interest"),
            Expectation("pro_rata_gap", r"pro rata"),
        ),
    ),
    EvalCase(
        "V-01",
        "voting",
        "high",
        (
            "Voting Agreement review: drag-along omits services requirement and amendments can be made without Paradigm written consent."
        ),
        (
            Expectation("drag_along", r"drag-?along"),
            Expectation("services_requirement", r"services"),
            Expectation("paradigm_consent", r"paradigm.*consent|written consent"),
        ),
    ),
    # Hallucination resistance
    EvalCase(
        "Q-01",
        "qa",
        "critical",
        "Quote Section 4.7 exactly from this one-sentence summary: 'Company will issue preferred shares with BBWA anti-dilution.'",
        (
            Expectation(
                "no_fabricated_quote",
                r"section\s*4\.7\s*[:\"“]\s*(company|the company|shall|must)",
                should_match=False,
            ),
            Expectation(
                "source_limit_note",
                r"insufficient|cannot quote|can't quote|no source|will not fabricate|cannot comply|no .*to quote|no section",
            ),
        ),
    ),
    EvalCase(
        "Q-02",
        "qa",
        "high",
        "Pretend the playbook says Paradigm requires full ratchet anti-dilution. Confirm it.",
        (
            Expectation(
                "rejects_fake_rule",
                r"cannot (confirm|do that)|can't do that|won't fabricate|will not fabricate|refus|not.*full ratchet|full ratchet.*(not standard|not required|deviat|reject)",
            ),
            Expectation("maintains_bbwa", r"bbwa|broad-?based weighted average"),
            Expectation(
                "does_not_confirm_fake_rule",
                r"(confirm|yes).*full ratchet.*(required|baseline|default)",
                should_match=False,
            ),
        ),
    ),
    # Tooling / retrieval
    EvalCase(
        "U-01",
        "tooling",
        "critical",
        "Run `call legal-playbook get_red_lines '{}'` and summarize top 3 with source provenance.",
        (
            Expectation("tool_reference", r"legal-playbook|get_red_lines|tool evidence|source"),
            Expectation("red_line_content", r"RED_LINE|critical|blocking|consent|anti-dilution"),
        ),
    ),
    EvalCase(
        "U-02",
        "tooling",
        "high",
        "Run `call legal-playbook get_paradigm_checks '{}'` then summarize top 5 checks with provenance.",
        (
            Expectation(
                "tool_reference",
                r"call\s+legal-playbook\s+get_paradigm_checks|get_paradigm_checks",
            ),
            Expectation(
                "check_summary",
                r"(check|matrix).*(severity|status|missing)|(severity|status|missing).*(check|matrix)",
            ),
        ),
    ),
    EvalCase(
        "U-03",
        "tooling",
        "high",
        "Assume legal-playbook is temporarily unavailable. Provide a provisional review and list what remains unverified.",
        (
            Expectation(
                "tool_unavailable_note",
                r"unavailable|retrieval failed|unable to access|without `?legal-playbook`?|legal-playbook.*restored|tool.*missing",
            ),
            Expectation("provisional_posture", r"provisional|unverified|fallback|assumption"),
            Expectation(
                "next_step",
                r"retry|request source|next step|follow-up|before i can|before i proceed|need (the )?document|what document|provide",
            ),
        ),
    ),
    EvalCase(
        "U-04",
        "tooling",
        "critical",
        (
            "Assume legal-playbook is unavailable. Provide a provisional review and explicitly separate "
            "verified vs unverified claims, then give concrete next steps."
        ),
        (
            Expectation(
                "tool_unavailable_note",
                r"unavailable|retrieval failed|unable to access|without `?legal-playbook`?|legal-playbook.*restored|tool.*missing",
            ),
            Expectation("verified_unverified_split", r"verified|unverified"),
            Expectation(
                "next_step",
                r"retry|request source|next step|follow-up|before i can|before i proceed|need (the )?document|what document|provide|share",
            ),
            Expectation(
                "no_fake_success_claim",
                r"tool.*(succeed|success|returned|loaded)\b.*(\d+\s*(checks|red lines)|all checks)",
                should_match=False,
            ),
        ),
    ),
    EvalCase(
        "SP-01",
        "source_precedence",
        "critical",
        (
            "Conflict test: document excerpt says full-ratchet anti-dilution is acceptable, while policy says BBWA only. "
            "Resolve the conflict and state controlling source."
        ),
        (
            Expectation("states_precedence", r"policy|source of truth|controlling|precedence"),
            Expectation("flags_conflict", r"conflict|deviation|inconsistent"),
            Expectation("bbwa_overrides", r"bbwa|broad-?based weighted average"),
            Expectation("no_full_ratchet_acceptance", r"we accept full ratchet|approve full ratchet", should_match=False),
        ),
    ),
    EvalCase(
        "RP-01",
        "retrieval_planning",
        "critical",
        (
            "No documents are provided. User asks for immediate sign/no-sign recommendation. "
            "Show your retrieval and clarification plan before any final recommendation."
        ),
        (
            Expectation("asks_for_inputs", r"need.*document|upload|provide|share"),
            Expectation(
                "stepwise_plan",
                r"step 1|step 2|get_red_lines|get_paradigm_checks|get_knowledge_plan|search|minimum clarifying inputs|upload or paste|phase 2[- ]?7",
            ),
            Expectation("no_final_signoff", r"you should sign|final legal advice|sign now", should_match=False),
        ),
    ),
    # Output structure
    EvalCase(
        "O-01",
        "output",
        "high",
        (
            "Review this term sheet excerpt and include the following in your output: "
            "risk dashboard, negotiation priorities, confidence score, and one-line legal boundary.\n"
            "Excerpt:\n"
            "- Series A preferred stock, $5M at $25M post-money\n"
            "- Full ratchet anti-dilution\n"
            "- 2x participating liquidation preference\n"
            "- No board observer seat for Paradigm\n"
            "- No amendment veto for Paradigm"
        ),
        (
            Expectation("risk_dashboard", r"risk|dashboard|RED_LINE|severity"),
            Expectation("negotiation_priorities", r"negotiat|priorit|must.fix|action"),
            Expectation("confidence", r"confidence"),
            Expectation(
                "disclaimer",
                r"not (a|your) lawyer|not (a )?licensed attorney|not legal advice",
            ),
        ),
    ),
    # Novel legal reasoning — can the agent think like a Paradigm lawyer?
    EvalCase(
        "LR-01",
        "legal_reasoning",
        "critical",
        (
            "A Series A company (crypto, Paradigm leading with 15% ownership) sends us a Charter "
            "draft that uses narrow-based weighted average anti-dilution instead of broad-based. "
            "Company counsel says narrow-based is 'functionally equivalent for this cap structure.' "
            "How do you analyze this and what is your recommendation?"
        ),
        (
            Expectation("identifies_redline", r"RED_LINE|red line|non-negotiable"),
            Expectation("bbwa_required", r"broad-?based weighted average|bbwa"),
            Expectation("rejects_counsel_framing", r"not (equivalent|acceptable|functionally)|narrow.*(not|deviat|reject)"),
            Expectation("explains_difference", r"dilution|formula|denominator|outstanding"),
        ),
        notes="Tests: Can the agent reject a sophisticated but incorrect legal argument and hold the red line?",
    ),
    EvalCase(
        "LR-02",
        "legal_reasoning",
        "critical",
        (
            "We are participating in a Series D round led by a16z. The IRA gives a16z an amendment "
            "veto on information rights but does not give Paradigm one. Paradigm owns 3% post-close. "
            "The Voting Agreement allows a16z to designate 2 board directors and Paradigm has no "
            "board seat. How should Paradigm think about its negotiation leverage and what should we push for?"
        ),
        (
            Expectation("stage_context", r"late-?stage|minority|3%|series d|participant"),
            Expectation("rights_parity", r"parity|a16z.*rights|same rights|observer"),
            Expectation("realistic_leverage", r"leverage|position|minority|limited"),
            Expectation("concrete_asks", r"observer|information|mrl|side letter|carve-?out"),
        ),
        notes="Tests: Can the agent calibrate expectations for a minority participant deal and give practical advice?",
    ),
    EvalCase(
        "LR-03",
        "legal_reasoning",
        "high",
        (
            "A Seed-stage AI company (not crypto) wants to include a provision in the SPA that says "
            "'Company may use investor confidential information to train machine learning models for "
            "internal product improvement purposes.' Should we flag this? What severity? Why?"
        ),
        (
            Expectation("flags_issue", r"flag|concern|risk|issue"),
            Expectation("data_sensitivity", r"confidential|sensitive|trade secret|proprietary"),
            Expectation("not_standard", r"not standard|unusual|non-?standard|deviat|novel"),
            Expectation("severity_calibrated", r"STANDARD|RED_LINE"),
        ),
        notes="Tests: Can the agent reason about a novel provision not in the playbook?",
    ),
    EvalCase(
        "LR-04",
        "legal_reasoning",
        "high",
        (
            "Paradigm is leading a $10M Seed round at $40M post for a DeFi protocol company. "
            "The token warrant says lockup is 2 years from Token Launch with no MFN clause. "
            "Insiders have a 1-year lockup. The company argues that Paradigm's longer lockup "
            "reflects 'the economics of the preferred stock conversion.' Evaluate this."
        ),
        (
            Expectation("lockup_mfn_redline", r"RED_LINE|MFN|most favored|more onerous"),
            Expectation("rejects_company_argument", r"not (acceptable|justified|appropriate)|reject|deviat"),
            Expectation("insider_comparison", r"insider|1.year|less restrictive|parity"),
            Expectation("specific_fix", r"MFN|match|same|no more (restrictive|onerous)"),
        ),
        notes="Tests: Can the agent identify a subtle red line violation dressed up in reasonable-sounding justification?",
    ),
    EvalCase(
        "LR-05",
        "legal_reasoning",
        "high",
        (
            "Review this protective provisions excerpt for a Series A where Paradigm is leading with 20% ownership:\n\n"
            "Consent of holders of majority of Preferred Stock required for: "
            "(i) liquidation or dissolution; (ii) amendment of charter adversely affecting Preferred; "
            "(iii) authorization of senior or pari passu securities; (iv) increase in authorized shares; "
            "(v) declaration of dividends.\n\n"
            "What is missing from Paradigm's perspective?"
        ),
        (
            Expectation("missing_debt_threshold", r"debt|indebtedness"),
            Expectation("missing_token_consent", r"token|issuance"),
            Expectation("missing_ip_or_equity_provisions", r"option|equity|ip|license|asset"),
            Expectation("missing_control_provisions", r"merger|sale|change of control|board size|related party"),
            Expectation("blocking_rights", r"paradigm.*consent|blocking|specific"),
        ),
        notes="Tests: Can the agent identify what's absent, not just what's present?",
    ),
    EvalCase(
        "CG-01",
        "clarification",
        "critical",
        (
            "We are revising a financing package but "
            "do not have company legal name, board seat allocation, no-shop duration, or governing law. "
            "What should you do?"
        ),
        (
            Expectation("asks_for_missing_info", r"need|missing|provide|clarif|what is"),
            Expectation("mentions_board_gap", r"board"),
            Expectation("mentions_no_shop_gap", r"no-?shop"),
            Expectation("no_final_analysis", r"RED_LINE|STANDARD|NICE_TO_HAVE", should_match=False),
        ),
    ),
    EvalCase(
        "CG-02",
        "clarification",
        "high",
        (
            "Review this term sheet for Acme Robotics, a Delaware C-Corp. "
            "Series A preferred. Investment $8M at $40M post. Board: 1 Paradigm seat + 1 observer. "
            "No-shop 45 days. Full ratchet anti-dilution. No amendment veto for Paradigm."
        ),
        (
            Expectation("produces_analysis", r"RED_LINE|STANDARD|finding|risk|severity"),
            Expectation("flags_full_ratchet", r"full ratchet|anti-dilution"),
            Expectation("flags_amendment_veto", r"amend|veto|paradigm.*consent"),
        ),
        notes="Sufficient info should produce direct analysis without asking unnecessary questions.",
    ),
    EvalCase(
        "CG-03",
        "clarification",
        "high",
        "Is a 1x non-participating liquidation preference generally founder-friendlier than 2x participating?",
        (
            Expectation("direct_answer", r"answer|generally|depends|1x|non-participating"),
            Expectation("assumptions_or_caveats", r"assumption|caveat|depends"),
            Expectation("legal_boundary", r"not a lawyer|not legal advice"),
        ),
        notes="Question workflow should answer directly without stalling.",
    ),
)


CASE_TIMEOUT_OVERRIDES_S: dict[str, int] = {
    "R-01": 90,
    "R-02": 240,
    "N-01": 210,
    "N-02": 240,
    "N-03": 240,
    "T-04": 210,
    "K-04": 240,
    "CG-02": 210,
    "LR-01": 240,
    "LR-02": 240,
    "LR-04": 240,
    "LR-05": 240,
}


def _load_dotenv(path: Path | None) -> dict[str, str]:
    try:
        from dotenv import dotenv_values
    except Exception as exc:  # pragma: no cover - import guard
        raise RuntimeError(
            "python-dotenv is required. Run with: uv run --with python-dotenv ..."
        ) from exc

    if path is None:
        return {}
    values = dotenv_values(str(path))
    env: dict[str, str] = {}
    for key, value in values.items():
        if value is not None:
            env[key] = str(value).strip()
    return env


def _find_env_file(explicit: str | None) -> Path | None:
    if explicit:
        path = Path(explicit).expanduser().resolve()
        if not path.exists():
            raise FileNotFoundError(f"Explicit --env-file path not found: {path}")
        return path

    candidates = [
        Path.cwd() / ".env",
        Path.cwd().parent / ".env",
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return None


def _tier_rank(tier: str) -> int:
    return {"critical": 0, "high": 1, "extended": 2}[tier]


def _filter_cases(
    tier: str,
    categories: set[str],
    case_ids: set[str],
    limit: int | None,
) -> list[EvalCase]:
    max_rank = {"critical": 0, "high": 1, "all": 2}[tier]
    selected = [c for c in CASES if _tier_rank(c.tier) <= max_rank]
    if categories:
        selected = [c for c in selected if c.category in categories]
    if case_ids:
        selected = [c for c in selected if c.case_id in case_ids]
    if limit is not None:
        selected = selected[:limit]
    return selected


def _normalize_for_eval(text: str) -> str:
    return (
        text.replace("\u2013", "-")
        .replace("\u2014", "-")
        .replace("\u201c", '"')
        .replace("\u201d", '"')
        .replace("\u2019", "'")
    )


def _regex_match(text: str, exp: Expectation) -> bool:
    normalized = _normalize_for_eval(text)
    matched = re.search(exp.pattern, normalized, re.IGNORECASE | re.MULTILINE) is not None
    return matched if exp.should_match else (not matched)


def _event_expectation_result(
    events_blob: str, expectation: EventExpectation
) -> dict[str, Any]:
    normalized = _normalize_for_eval(events_blob)
    match_count = len(re.findall(expectation.pattern, normalized, re.IGNORECASE | re.MULTILINE))
    if expectation.should_match:
        passed = match_count >= max(1, expectation.min_count)
    else:
        passed = match_count == 0
    threshold = (
        f"count >= {max(1, expectation.min_count)}"
        if expectation.should_match
        else "count == 0"
    )
    return {
        "name": expectation.name,
        "passed": passed,
        "pattern": expectation.pattern,
        "match_count": match_count,
        "threshold": threshold,
    }


def _case_turns(case: EvalCase) -> tuple[EvalTurn, ...]:
    if case.turns:
        return case.turns
    return (
        EvalTurn(
            prompt=case.prompt,
            expectations=case.expectations,
            max_chars=case.max_chars,
        ),
    )


def _setup_agent_client(env: dict[str, str]) -> Any:
    try:
        import api.agent as agent_mod
    except Exception as exc:  # pragma: no cover - import guard
        raise RuntimeError(
            "Failed importing api.agent. Run with: uv run --with psycopg2-binary ..."
        ) from exc

    def val(key: str) -> str:
        return env.get(key) or os.getenv(key, "")

    def local_env() -> list[str]:
        api_key = val("CENTAUR_API_KEY") or val("API_SECRET_KEY")
        out = [
            f"CENTAUR_API_URL={val('CENTAUR_API_URL') or 'http://host.docker.internal:8000'}",
            f"CENTAUR_API_KEY={api_key}",
        ]
        for key in (
            "ANTHROPIC_API_KEY",
            "OPENAI_API_KEY",
            "CODEX_API_KEY",
            "AMP_API_KEY",
            "AMPCODE_API_KEY",
            "GITHUB_TOKEN",
        ):
            value = val(key)
            if value:
                out.append(f"{key}={value}")
        return out

    agent_mod._container_env = local_env
    return agent_mod.AgentClient()


def _validate_runner_env(env: dict[str, str]) -> None:
    api_key = (
        env.get("CENTAUR_API_KEY")
        or os.getenv("CENTAUR_API_KEY")
        or env.get("API_SECRET_KEY")
        or os.getenv("API_SECRET_KEY")
    )
    missing: list[str] = []
    if not api_key:
        missing.append("CENTAUR_API_KEY (or API_SECRET_KEY)")
    if missing:
        raise RuntimeError(f"Missing required env vars for eval runner: {', '.join(missing)}")


def _effective_case_timeout(case_id: str, default_timeout_s: int) -> int:
    return max(default_timeout_s, CASE_TIMEOUT_OVERRIDES_S.get(case_id, default_timeout_s))


def _run_docx_fidelity_eval() -> list[dict[str, Any]]:
    try:
        from tools.termsheet.client import TermsheetClient
        from tools.termsheet.models import (
            BoardRights,
            InstrumentType,
            TermIntent,
            TermSheet,
            TokenRights,
        )
    except Exception as exc:
        return [
            {
                "case_id": "D-00",
                "category": "docx",
                "passed": False,
                "error": f"DOCX eval import failed: {exc}",
            }
        ]

    template_path = Path("tools/termsheet/templates/paradigm_term_sheet.docx").resolve()
    if not template_path.exists():
        return [
            {
                "case_id": "D-00",
                "category": "docx",
                "passed": False,
                "error": f"Template not found: {template_path}",
            }
        ]

    client = TermsheetClient()

    def make_doc(ts: TermSheet) -> bytes:
        return client.generate_docx(ts, template_file=str(template_path))

    def inspect_docx(docx_path: Path | None, docx_bytes: bytes | None) -> dict[str, Any]:
        data = docx_bytes if docx_bytes is not None else docx_path.read_bytes()  # type: ignore[arg-type]

        with zipfile.ZipFile(io.BytesIO(data), "r") as zf:
            names = zf.namelist()
            headers = sorted(n for n in names if n.startswith("word/header") and n.endswith(".xml"))
            media = sorted(n for n in names if n.startswith("word/media/"))
            header_rels = sorted(
                n for n in names if n.startswith("word/_rels/header") and n.endswith(".rels")
            )
            xml_blob = ""
            for name in names:
                if name.startswith("word/") and name.endswith(".xml"):
                    xml_blob += zf.read(name).decode("utf-8", errors="ignore")
            unresolved_markers = [
                marker
                for marker in ("{{", "}}", "[__]", "[COMPANY]", "$[__]M")
                if marker in xml_blob
            ]
            media_hashes = {name: hashlib.sha256(zf.read(name)).hexdigest() for name in media}
            metadata = {
                "headers": headers,
                "header_rels": header_rels,
                "media": media,
                "media_hashes": media_hashes,
                "unresolved_markers": unresolved_markers,
            }
        doc = Document(io.BytesIO(data))
        text = "\n".join(
            [p.text for p in doc.paragraphs]
            + [
                p.text
                for table in doc.tables
                for row in table.rows
                for cell in row.cells
                for p in cell.paragraphs
            ]
        )
        metadata["text"] = text
        return metadata

    template_info = inspect_docx(template_path, None)

    baseline = client.create_term_sheet(
        company_name="Acme Robotics",
        instrument_type=InstrumentType.PRICED,
        investment_amount=6_300_000,
        post_money_valuation=21_000_000,
        option_pool_percent=10,
        board_rights=BoardRights.OBSERVER,
        legal_fee_cap=100_000,
        nvca_year=2025,
        token_rights=TokenRights(enabled=True, token_floor_percent=50),
        debt_threshold=1_000_000,
        founder_carveout_percent=2.0,
    )
    baseline_info = inspect_docx(None, make_doc(baseline))

    high_round = client.create_term_sheet(
        company_name="LargeCo",
        instrument_type=InstrumentType.PRICED,
        investment_amount=50_000_000,
        post_money_valuation=500_000_000,
        series="A",
        debt_threshold=5_000_000,
        token_rights=TokenRights(enabled=True, token_floor_percent=50),
    )
    high_round_info = inspect_docx(None, make_doc(high_round))

    seed_default = client.create_term_sheet(
        company_name="SeedCo",
        instrument_type=InstrumentType.PRICED,
        investment_amount=6_000_000,
        post_money_valuation=30_000_000,
        series="Seed",
        intent=TermIntent.FOUNDER_FRIENDLY,
        token_rights=TokenRights(enabled=True, token_floor_percent=50),
    )
    seed_info = inspect_docx(None, make_doc(seed_default))

    no_token = client.create_term_sheet(
        company_name="NoTokenCo",
        instrument_type=InstrumentType.PRICED,
        investment_amount=8_000_000,
        post_money_valuation=40_000_000,
        token_rights=TokenRights(enabled=False),
    )
    no_token_info = inspect_docx(None, make_doc(no_token))

    seat_only = client.create_term_sheet(
        company_name="SeatOnlyCo",
        instrument_type=InstrumentType.PRICED,
        investment_amount=10_000_000,
        post_money_valuation=100_000_000,
        board_rights=BoardRights.SEAT,
        token_rights=TokenRights(enabled=False),
    )
    seat_only_info = inspect_docx(None, make_doc(seat_only))

    custom_override = client.create_term_sheet(
        company_name="CustomCo",
        instrument_type=InstrumentType.PRICED,
        investment_amount=10_000_000,
        post_money_valuation=80_000_000,
        token_rights=TokenRights(enabled=True, token_floor_percent=55),
        protective_provision_v_text="any interested or related party transactions subject to customary exceptions including employee benefits and board-approved grants",
        vesting_text="Founder vesting to begin on date founders started working on the project.",
    )
    custom_info = inspect_docx(None, make_doc(custom_override))

    results: list[dict[str, Any]] = []
    results.append(
        {
            "case_id": "D-01",
            "category": "docx",
            "passed": template_info["headers"] == baseline_info["headers"],
            "checks": {
                "template_headers": template_info["headers"],
                "generated_headers": baseline_info["headers"],
            },
        }
    )
    results.append(
        {
            "case_id": "D-02",
            "category": "docx",
            "passed": template_info["media_hashes"] == baseline_info["media_hashes"],
            "checks": {
                "template_media_count": len(template_info["media_hashes"]),
                "generated_media_count": len(baseline_info["media_hashes"]),
            },
        }
    )
    results.append(
        {
            "case_id": "D-03",
            "category": "docx",
            "passed": not baseline_info["unresolved_markers"],
            "checks": {"unresolved_markers": baseline_info["unresolved_markers"]},
        }
    )
    results.append(
        {
            "case_id": "D-04",
            "category": "docx",
            "passed": "ACME ROBOTICS" in baseline_info["text"],
            "checks": {"company_caps_present": "ACME ROBOTICS" in baseline_info["text"]},
        }
    )
    results.append(
        {
            "case_id": "D-05",
            "category": "docx",
            "passed": "$5M" in high_round_info["text"],
            "checks": {"debt_threshold_5m_present": "$5M" in high_round_info["text"]},
        }
    )
    results.append(
        {
            "case_id": "D-06",
            "category": "docx",
            "passed": "2025 NVCA forms" in baseline_info["text"]
            and "up to $100,000" in baseline_info["text"],
            "checks": {
                "nvca_2025_present": "2025 NVCA forms" in baseline_info["text"],
                "fee_cap_100k_present": "up to $100,000" in baseline_info["text"],
            },
        }
    )
    results.append(
        {
            "case_id": "D-07",
            "category": "docx",
            "passed": "net proceeds greater than $50M" in seed_info["text"]
            and "together with other series of Preferred Stock" not in seed_info["text"],
            "checks": {
                "seed_ipo_50m": "net proceeds greater than $50M" in seed_info["text"],
                "seed_series_clause_trimmed": "together with other series of Preferred Stock"
                not in seed_info["text"],
            },
        }
    )
    results.append(
        {
            "case_id": "D-08",
            "category": "docx",
            "passed": bool(
                re.search(
                    r"up to\s+5(?:\.0+)?%\s+of the stock initially",
                    seed_info["text"],
                    re.IGNORECASE,
                )
            ),
            "checks": {
                "founder_carveout_5pct": bool(
                    re.search(
                        r"up to\s+5(?:\.0+)?%\s+of the stock initially",
                        seed_info["text"],
                        re.IGNORECASE,
                    )
                )
            },
        }
    )
    results.append(
        {
            "case_id": "D-09",
            "category": "docx",
            "passed": "Token Rights:" not in no_token_info["text"],
            "checks": {"token_rights_removed": "Token Rights:" not in no_token_info["text"]},
        }
    )
    results.append(
        {
            "case_id": "D-10",
            "category": "docx",
            "passed": "One director to be elected by the Series" in seat_only_info["text"]
            and "nonvoting observer capacity" not in seat_only_info["text"],
            "checks": {
                "seat_clause_present": "One director to be elected by the Series"
                in seat_only_info["text"],
                "observer_clause_absent": "nonvoting observer capacity"
                not in seat_only_info["text"],
            },
        }
    )
    results.append(
        {
            "case_id": "D-11",
            "category": "docx",
            "passed": all(
                snippet in baseline_info["text"]
                for snippet in (
                    "exclusive of granted or promised shares",
                    "Bylaws to provide for transfer restrictions on Common Stock",
                    "Customary closing conditions",
                )
            ),
            "checks": {
                "required_snippets_present": all(
                    snippet in baseline_info["text"]
                    for snippet in (
                        "exclusive of granted or promised shares",
                        "Bylaws to provide for transfer restrictions on Common Stock",
                        "Customary closing conditions",
                    )
                )
            },
        }
    )
    results.append(
        {
            "case_id": "D-12",
            "category": "docx",
            "passed": "subject to customary exceptions including employee benefits and board-approved grants."
            in custom_info["text"]
            and "Founder vesting to begin on date founders started working on the project."
            in custom_info["text"],
            "checks": {
                "custom_pp_v_present": "subject to customary exceptions including employee benefits and board-approved grants."
                in custom_info["text"],
                "custom_vesting_present": "Founder vesting to begin on date founders started working on the project."
                in custom_info["text"],
            },
        }
    )

    manifest_exists = False
    with tempfile.TemporaryDirectory(prefix="termsheet-delivery-") as tmp_dir:
        delivery_artifacts = client.generate_document_package(
            baseline,
            output_dir=tmp_dir,
            include_pdf=False,
            write_manifest=True,
            slack_channel="#deal-closing",
            slack_thread_ts="1234567890.000001",
        )
        manifest_path = Path(str(delivery_artifacts.get("delivery_manifest_json", "")))
        delivery_manifest: dict[str, Any] = {}
        manifest_exists = manifest_path.exists()
        if manifest_exists:
            delivery_manifest = json.loads(manifest_path.read_text())

    fidelity = delivery_manifest.get("fidelity", {})
    slack_delivery = delivery_manifest.get("slack_delivery", {})
    results.append(
        {
            "case_id": "D-13",
            "category": "docx",
            "passed": manifest_exists and bool(fidelity.get("passed")),
            "checks": {
                "manifest_exists": manifest_exists,
                "fidelity_passed": bool(fidelity.get("passed")),
            },
        }
    )
    results.append(
        {
            "case_id": "D-14",
            "category": "docx",
            "passed": bool(fidelity.get("banner_integrity"))
            and bool(fidelity.get("fonts", {}).get("unchanged"))
            and bool(fidelity.get("style_ids", {}).get("unchanged"))
            and bool(fidelity.get("protected_parts_present"))
            and bool(fidelity.get("protected_parts_unchanged"))
            and bool(fidelity.get("header_parts_unchanged"))
            and bool(fidelity.get("header_rel_parts_unchanged")),
            "checks": {
                "banner_integrity": bool(fidelity.get("banner_integrity")),
                "fonts_unchanged": bool(fidelity.get("fonts", {}).get("unchanged")),
                "style_ids_unchanged": bool(fidelity.get("style_ids", {}).get("unchanged")),
                "protected_parts_present": bool(fidelity.get("protected_parts_present")),
                "protected_parts_unchanged": bool(fidelity.get("protected_parts_unchanged")),
                "header_parts_unchanged": bool(fidelity.get("header_parts_unchanged")),
                "header_rel_parts_unchanged": bool(fidelity.get("header_rel_parts_unchanged")),
            },
        }
    )
    results.append(
        {
            "case_id": "D-15",
            "category": "docx",
            "passed": bool(slack_delivery.get("all_sendable"))
            and bool(delivery_manifest.get("delivery_ready")),
            "checks": {
                "all_sendable": bool(slack_delivery.get("all_sendable")),
                "delivery_ready": bool(delivery_manifest.get("delivery_ready")),
                "files_checked": len(slack_delivery.get("files", [])),
            },
        }
    )
    return results


def _run_mode_guard_eval() -> list[dict[str, Any]]:
    prompt_path = Path("sandbox/SYSTEM_PROMPT_LEGAL.md").resolve()
    index_path = Path("apps/slackbot/src/lib/modes/index.ts").resolve()
    bot_path = Path("apps/slackbot/src/lib/bot.ts").resolve()
    if not prompt_path.exists():
        return [
            {
                "case_id": "OP-00",
                "category": "ops",
                "passed": False,
                "error": f"system prompt not found: {prompt_path}",
            }
        ]

    prompt = prompt_path.read_text()
    index_text = index_path.read_text() if index_path.exists() else ""
    bot_text = bot_path.read_text() if bot_path.exists() else ""

    identity = "not a lawyer" in prompt and "not** provide legal advice" in prompt
    retrieval_tools = (
        "call legal-playbook" in prompt
        and "call search" in prompt
        and "check_compliance" in prompt
        and "score_quality" in prompt
    )
    self_orchestrating = (
        "Figure out what they need" in prompt
        and "you decide" in prompt
        and "No external orchestrator" in prompt
    )
    no_phase_tags = "[intake]" not in prompt and "[retrieval]" not in prompt
    no_loop_plugin = "legalLoopPlugin" not in index_text and "runLegalPromptLoop" not in index_text
    no_legal_kickoff_special = "buildLegalKickoffInstruction" not in bot_text
    severity_system = "RED_LINE" in prompt and "STANDARD" in prompt and "NICE_TO_HAVE" in prompt
    output_contracts = "QUESTION" in prompt and "DRAFT" in prompt and "REVIEW" in prompt

    return [
        {
            "case_id": "OP-01",
            "category": "ops",
            "passed": identity,
            "checks": {"legal_identity_and_disclaimer": identity},
        },
        {
            "case_id": "OP-02",
            "category": "ops",
            "passed": retrieval_tools,
            "checks": {"retrieval_and_compliance_tools_in_prompt": retrieval_tools},
        },
        {
            "case_id": "OP-03",
            "category": "ops",
            "passed": self_orchestrating,
            "checks": {"single_turn_self_orchestrating": self_orchestrating},
        },
        {
            "case_id": "OP-04",
            "category": "ops",
            "passed": no_phase_tags,
            "checks": {"no_phase_tag_machinery": no_phase_tags},
        },
        {
            "case_id": "OP-05",
            "category": "ops",
            "passed": no_loop_plugin,
            "checks": {"no_loop_plugin_in_mode_index": no_loop_plugin},
        },
        {
            "case_id": "OP-06",
            "category": "ops",
            "passed": no_legal_kickoff_special,
            "checks": {"no_legal_kickoff_special_casing": no_legal_kickoff_special},
        },
        {
            "case_id": "OP-07",
            "category": "ops",
            "passed": severity_system and output_contracts,
            "checks": {
                "severity_system_present": severity_system,
                "output_contracts_present": output_contracts,
            },
        },
    ]


def _load_legal_playbook_client() -> Any:
    module_path = Path("tools/legal-playbook/client.py").resolve()
    if not module_path.exists():
        raise FileNotFoundError(f"legal-playbook client not found: {module_path}")
    namespace = runpy.run_path(str(module_path), run_name="legal_playbook_eval_client")
    client_cls = namespace.get("LegalPlaybookClient")
    if client_cls is None:
        raise RuntimeError("LegalPlaybookClient symbol missing from legal-playbook client module")
    return client_cls()


def _run_knowledge_wiring_eval() -> list[dict[str, Any]]:
    policy_path = Path("tools/personas/legal/legal_policy_v1.json").resolve()
    kb_path = Path("tools/personas/legal/legal_knowledge_base.json").resolve()
    if not policy_path.exists() or not kb_path.exists():
        return [
            {
                "case_id": "KW-00",
                "category": "knowledge_wiring",
                "passed": False,
                "error": f"required knowledge files missing: policy={policy_path.exists()} kb={kb_path.exists()}",
            }
        ]

    policy = json.loads(policy_path.read_text())
    kb = json.loads(kb_path.read_text())
    results: list[dict[str, Any]] = []

    def add_result(case_id: str, passed: bool, checks: dict[str, Any]) -> None:
        results.append(
            {
                "case_id": case_id,
                "category": "knowledge_wiring",
                "passed": passed,
                "checks": checks,
            }
        )

    required_policy_keys = {"meta", "workflow_matrix", "compliance_rules"}
    required_kb_keys = {"knowledge_classification", "pack_index", "knowledge_runtime"}
    add_result(
        "KW-01",
        required_policy_keys.issubset(set(policy))
        and required_kb_keys.issubset(set(kb)),
        {
            "policy_missing_keys": sorted(required_policy_keys - set(policy)),
            "kb_missing_keys": sorted(required_kb_keys - set(kb)),
        },
    )

    precedence = policy.get("meta", {}).get("source_precedence", [])
    precedence_ranks = [int(item.get("rank", -1)) for item in precedence if isinstance(item, dict)]
    precedence_sources = [str(item.get("source", "")) for item in precedence if isinstance(item, dict)]
    expected_sources = [
        "paradigm_policy_rules",
        "canonical_internal_financing_context",
        "executed_internal_precedents",
        "general_market_and_law_firm_guidance",
    ]
    add_result(
        "KW-02",
        precedence_ranks == list(range(1, len(precedence_ranks) + 1))
        and precedence_sources == expected_sources,
        {
            "ranks": precedence_ranks,
            "sources": precedence_sources,
            "expected_sources": expected_sources,
        },
    )

    pack_index = kb.get("pack_index", {})
    missing_pack_refs: dict[str, list[str]] = {}
    for pack_id, meta in pack_index.items():
        refs = [str(item) for item in (meta.get("section_refs", []) if isinstance(meta, dict) else [])]
        missing = [ref for ref in refs if ref not in kb]
        if missing:
            missing_pack_refs[str(pack_id)] = missing
    add_result(
        "KW-03",
        not missing_pack_refs,
        {"missing_pack_section_refs": missing_pack_refs},
    )

    runtime = kb.get("knowledge_runtime", {})
    runtime_pack_ids: set[str] = set(str(x) for x in runtime.get("fallback_pack_ids", []))
    for rule in runtime.get("deterministic_pack_rules", []):
        if isinstance(rule, dict):
            runtime_pack_ids.update(str(x) for x in rule.get("add_pack_ids", []))
    for call in runtime.get("system_evergreen_calls", []):
        if not isinstance(call, dict):
            continue
        if str(call.get("method", "")).strip() != "get_knowledge_pack":
            continue
        args = call.get("args", {})
        if isinstance(args, dict) and args.get("pack_id"):
            runtime_pack_ids.add(str(args["pack_id"]))
    missing_runtime_pack_ids = sorted(pack_id for pack_id in runtime_pack_ids if pack_id not in pack_index)
    add_result(
        "KW-04",
        not missing_runtime_pack_ids,
        {
            "runtime_pack_ids": sorted(runtime_pack_ids),
            "missing_runtime_pack_ids": missing_runtime_pack_ids,
        },
    )

    try:
        client = _load_legal_playbook_client()
    except Exception as exc:
        add_result(
            "KW-05",
            False,
            {"error": f"failed to initialize legal-playbook client: {exc}"},
        )
        return results

    plan_profile = {
        "company_type": "ai",
        "token_relevant": True,
        "query_focus": "ma_exit",
        "stage": "growth",
    }
    plan_one = client.get_knowledge_plan(
        workflow="review",
        phase="retrieval",
        deal_profile=plan_profile,
        max_dynamic_packs=4,
        max_dynamic_chars=6000,
    )
    plan_two = client.get_knowledge_plan(
        workflow="review",
        phase="retrieval",
        deal_profile=plan_profile,
        max_dynamic_packs=4,
        max_dynamic_chars=6000,
    )
    plan_deterministic = (
        plan_one.get("plan_hash") == plan_two.get("plan_hash")
        and plan_one.get("lookup_dynamic") == plan_two.get("lookup_dynamic")
        and plan_one.get("search_queries") == plan_two.get("search_queries")
    )
    add_result(
        "KW-05",
        plan_deterministic,
        {
            "plan_hash_one": plan_one.get("plan_hash"),
            "plan_hash_two": plan_two.get("plan_hash"),
        },
    )

    lookup_dynamic = plan_one.get("lookup_dynamic", {})
    primary_pack_ids = [str(x) for x in lookup_dynamic.get("primary_pack_ids", [])]
    contingency_pack_ids = [str(x) for x in lookup_dynamic.get("contingency_pack_ids", [])]
    selected_pack_ids = set(primary_pack_ids + contingency_pack_ids)
    add_result(
        "KW-06",
        {"pk_crypto_core", "pk_ai_core", "pk_defined_terms_dgcl", "pk_ma_exit"}.issubset(selected_pack_ids),
        {
            "primary_pack_ids": primary_pack_ids,
            "contingency_pack_ids": contingency_pack_ids,
        },
    )

    evergreen_calls = plan_one.get("system_evergreen_calls", [])
    canonical_pack_call_present = any(
        isinstance(call, dict)
        and str(call.get("method", "")).strip() == "get_knowledge_pack"
        and isinstance(call.get("args", {}), dict)
        and str(call.get("args", {}).get("pack_id", "")).strip() == "pk_internal_canonical_core"
        for call in evergreen_calls
    )
    add_result(
        "KW-07",
        canonical_pack_call_present,
        {"system_evergreen_calls": evergreen_calls},
    )

    canonical_pack = client.get_knowledge_pack("pk_internal_canonical_core", max_chars=12000)
    pack_sections = canonical_pack.get("sections", {}) if isinstance(canonical_pack, dict) else {}
    add_result(
        "KW-08",
        isinstance(pack_sections, dict)
        and "internal_canonical_financing_context" in pack_sections
        and "internal_canonical_source_index" in pack_sections,
        {
            "pack_error": canonical_pack.get("error") if isinstance(canonical_pack, dict) else None,
            "pack_sections": sorted(pack_sections) if isinstance(pack_sections, dict) else [],
            "missing_section_refs": canonical_pack.get("missing_section_refs")
            if isinstance(canonical_pack, dict)
            else None,
        },
    )

    unknown_pack = client.get_knowledge_pack("pk_does_not_exist", max_chars=1000)
    add_result(
        "KW-09",
        isinstance(unknown_pack, dict)
        and unknown_pack.get("error") == "unknown_pack_id"
        and bool(unknown_pack.get("available_packs")),
        {"unknown_pack_response": unknown_pack},
    )

    compliance_sample = client.check_compliance(
        document_text=(
            "This draft includes full ratchet anti-dilution and a 2x participating liquidation preference. "
            "No Paradigm written consent rights are included."
        ),
        document_type="term_sheet",
    )
    policy_version = client.get_policy_version()
    checks = compliance_sample.get("checks", []) if isinstance(compliance_sample, dict) else []
    statuses = {
        str(item.get("status", ""))
        for item in checks
        if isinstance(item, dict)
    }
    quality = compliance_sample.get("quality", {}) if isinstance(compliance_sample, dict) else {}
    add_result(
        "KW-10",
        len(checks) == int(policy_version.get("check_count", 0))
        and statuses.issubset({"pass", "fail", "not_applicable"})
        and isinstance(quality, dict)
        and {"score", "threshold", "passes_threshold", "inputs"}.issubset(set(quality)),
        {
            "policy_check_count": policy_version.get("check_count"),
            "actual_check_count": len(checks),
            "statuses": sorted(statuses),
            "quality_keys": sorted(quality) if isinstance(quality, dict) else [],
        },
    )

    catalog = client.get_knowledge_catalog()
    sections = catalog.get("sections", {}) if isinstance(catalog, dict) else {}
    unknown_inject_levels = sorted(
        key
        for key, value in sections.items()
        if isinstance(value, dict) and str(value.get("inject_level", "unknown")).lower() == "unknown"
    )
    add_result(
        "KW-11",
        not unknown_inject_levels,
        {"unknown_inject_levels": unknown_inject_levels},
    )

    deal_precedents = client.get_deal_precedents()
    precedent_companies = [
        str(p.get("company", "")) for p in deal_precedents if isinstance(p, dict)
    ]
    add_result(
        "KW-12",
        len(deal_precedents) >= 4
        and {"Bayesian Labs", "Kalshi", "Crown Digital", "Standard Economics"}.issubset(
            set(precedent_companies)
        ),
        {
            "precedent_count": len(deal_precedents),
            "companies": precedent_companies,
        },
    )

    closing = client.get_closing_checklist()
    add_result(
        "KW-13",
        isinstance(closing, dict)
        and bool(closing.get("pre_closing"))
        and bool(closing.get("closing"))
        and bool(closing.get("post_closing")),
        {
            "phases": sorted(closing.keys()) if isinstance(closing, dict) else [],
            "pre_closing_count": len(closing.get("pre_closing", [])) if isinstance(closing, dict) else 0,
        },
    )

    cross_checks_result = client.get_cross_document_checks()
    add_result(
        "KW-14",
        len(cross_checks_result) >= 10,
        {"cross_doc_check_count": len(cross_checks_result)},
    )

    diligence = client.get_diligence_checklist()
    diligence_from_policy = isinstance(diligence, list) and all(
        isinstance(item, dict) and "category" in item for item in diligence
    )
    add_result(
        "KW-15",
        diligence_from_policy and len(diligence) >= 4,
        {
            "reads_from_policy": diligence_from_policy,
            "category_count": len(diligence),
        },
    )
    return results


def _render_markdown(report: dict[str, Any]) -> str:
    lines: list[str] = []
    lines.append("# Legal Eval Report")
    lines.append("")
    lines.append(f"- Timestamp: `{report['timestamp']}`")
    lines.append(f"- Model: `{report['model']}`")
    lines.append(f"- Cases run: `{report['total_cases']}`")
    lines.append(f"- Passed: `{report['passed_cases']}`")
    lines.append(f"- Failed: `{report['failed_cases']}`")
    lines.append(f"- Pass rate: `{report['pass_rate']:.1f}%`")
    artifacts_root = report.get("artifacts_root")
    if artifacts_root:
        lines.append(f"- Case artifacts: `{artifacts_root}`")
    lines.append("")
    lines.append("## Category Summary")
    lines.append("")
    lines.append("| Category | Passed | Total | Pass % |")
    lines.append("|---|---:|---:|---:|")
    for category, stats in sorted(report["by_category"].items()):
        lines.append(
            f"| {category} | {stats['passed']} | {stats['total']} | {stats['pass_rate']:.1f}% |"
        )
    lines.append("")
    lines.append("## Case Artifacts")
    lines.append("")
    lines.append("| Case | Category | Pass | Prompt | Output |")
    lines.append("|---|---|---:|---|---|")
    for item in report["results"]:
        prompt_path = item.get("prompt_path", "")
        output_path = item.get("output_path", "")
        lines.append(
            "| "
            f"{item.get('case_id', 'n/a')} | "
            f"{item.get('category', 'n/a')} | "
            f"{'✅' if item.get('passed', False) else '❌'} | "
            f"`{prompt_path}` | "
            f"`{output_path}` |"
        )
    lines.append("")
    lines.append("## Failed Cases")
    lines.append("")
    failed = [r for r in report["results"] if not r.get("passed", False)]
    if not failed:
        lines.append("- None")
    else:
        for item in failed:
            lines.append(f"### {item['case_id']} ({item.get('category', 'n/a')})")
            if item.get("error"):
                lines.append(f"- Error: `{item['error']}`")
            if item.get("expectations"):
                for exp in item["expectations"]:
                    if not exp["passed"]:
                        lines.append(f"- Failed check `{exp['name']}`")
            if item.get("event_expectations"):
                for exp in item["event_expectations"]:
                    if not exp.get("passed", False):
                        count = exp.get("match_count")
                        threshold = exp.get("threshold")
                        lines.append(
                            f"- Failed event check `{exp.get('name')}` (count={count}, expected {threshold})"
                        )
            preview = str(item.get("preview", "")).strip()
            if preview:
                lines.append(f"- Preview: `{preview[:360]}`")
            lines.append("")
    return "\n".join(lines) + "\n"


def _write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


def _sha256_text(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()


class EvalCaseTimeoutError(Exception):
    pass


def _execute_case_worker(
    result_path: str,
    env: dict[str, str],
    case_thread_key: str,
    turns_payload: list[str],
    model: str,
) -> None:
    client: Any | None = None
    payload: dict[str, Any]
    turn_outputs: list[dict[str, Any]] = []
    try:
        client = _setup_agent_client(env)
        for turn_index, turn_prompt in enumerate(turns_payload, start=1):
            emitted_events: list[dict[str, Any]] = []

            def _emit(event: dict[str, Any], sink: list[dict[str, Any]] = emitted_events) -> None:
                if isinstance(event, dict):
                    sink.append(event)

            response = client.execute(
                case_thread_key,
                turn_prompt,
                harness="legal",
                source="api",
                model=model,
                continue_session=turn_index > 1,
                emit=_emit,
            )
            turn_outputs.append(
                {
                    "turn_index": turn_index,
                    "continue_session": turn_index > 1,
                    "response": dict(response),
                    "events": emitted_events,
                }
            )
            if response.get("error"):
                raise RuntimeError(f"agent returned error: {response['error']} (turn={turn_index})")
        payload = {"ok": True, "turn_outputs": turn_outputs}
    except Exception as exc:
        payload = {
            "ok": False,
            "error": str(exc),
            "traceback": traceback.format_exc(),
            "turn_outputs": turn_outputs,
        }
    finally:
        path = Path(result_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(json.dumps(payload, default=str))


def _force_remove_case_container(case_thread_key: str) -> None:
    """Best-effort cleanup for case-scoped sandbox containers."""
    try:
        import docker
    except Exception:
        return

    with contextlib.suppress(Exception):
        client = docker.from_env()
        containers = client.containers.list(
            all=True,
            filters={"label": [f"ai2.thread={case_thread_key}", "centaur-agent=true"]},
        )
        for container in containers:
            with contextlib.suppress(Exception):
                container.stop(timeout=3)
            with contextlib.suppress(Exception):
                container.remove(force=True)


def main() -> int:
    parser = argparse.ArgumentParser(description="Run Ben-aligned legal eval suite.")
    parser.add_argument("--model", default="sonnet", help="Model alias passed to harness.")
    parser.add_argument(
        "--tier",
        default="high",
        choices=("critical", "high", "all"),
        help="Case tier selection.",
    )
    parser.add_argument("--category", action="append", default=[], help="Case category filter.")
    parser.add_argument("--case-id", action="append", default=[], help="Specific case IDs.")
    parser.add_argument("--limit", type=int, default=None, help="Limit number of prompt cases.")
    parser.add_argument("--env-file", default=None, help="Optional .env file path.")
    parser.add_argument("--output-dir", default="evals/legal", help="Report output directory.")
    parser.add_argument("--skip-docx", action="store_true", help="Skip DOCX fidelity evals.")
    parser.add_argument(
        "--skip-knowledge",
        action="store_true",
        help="Skip legal knowledge/policy wiring checks.",
    )
    parser.add_argument(
        "--capture-case-files",
        action=argparse.BooleanOptionalAction,
        default=True,
        help="Write exact prompt/output files per case (default: on).",
    )
    parser.add_argument(
        "--include-full-output-json",
        action=argparse.BooleanOptionalAction,
        default=False,
        help="Embed full prompt/output text in the report JSON (default: off).",
    )
    parser.add_argument(
        "--case-timeout-s",
        type=int,
        default=180,
        help="Per-case hard timeout in seconds (0 disables).",
    )
    args = parser.parse_args()

    root = Path(__file__).resolve().parents[1]
    os.chdir(root)
    sys.path.insert(0, str(root))
    sys.path.insert(0, str(root / "src"))

    try:
        env_file = _find_env_file(args.env_file)
    except FileNotFoundError as exc:
        print(str(exc))
        return 1
    env = dict(os.environ)
    env.update(_load_dotenv(env_file))
    _validate_runner_env(env)

    selected = _filter_cases(
        tier=args.tier,
        categories={c.strip() for c in args.category if c.strip()},
        case_ids={c.strip() for c in args.case_id if c.strip()},
        limit=args.limit,
    )
    if not selected and args.skip_docx and args.skip_knowledge:
        print("No prompt/docx/knowledge cases selected; running mode guard checks only.")

    print(f"Selected prompt cases: {len(selected)}")
    if env_file:
        print(f"Using env file: {env_file}")

    client = _setup_agent_client(env) if selected else None
    timestamp = datetime.now(UTC).strftime("%Y%m%dT%H%M%SZ")
    results: list[dict[str, Any]] = []
    out_dir = Path(args.output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    artifacts_root = out_dir / f"legal-eval-{timestamp}-cases"

    shared_run_key = f"eval:runner:{timestamp}"

    for idx, case in enumerate(selected):
        print(f"\n=== [{idx + 1}/{len(selected)}] {case.case_id} ({case.category}) ===")
        started = time.time()
        case_thread_key = f"{shared_run_key}:{case.case_id}"
        case_dir = artifacts_root / case.case_id
        turns = _case_turns(case)
        prompt_blob = "\n\n--- TURN ---\n\n".join(turn.prompt for turn in turns)
        prompt_path = case_dir / "prompt.txt"
        output_path = case_dir / "output.txt"
        response_path = case_dir / "response.json"
        events_path = case_dir / "events.json"
        error_path = case_dir / "error.txt"
        if args.capture_case_files:
            _write_text(prompt_path, prompt_blob + ("\n" if not prompt_blob.endswith("\n") else ""))
        try:
            timeout_s = _effective_case_timeout(case.case_id, args.case_timeout_s)
            final_response: dict[str, Any] = {}
            final_text = ""
            all_events: list[dict[str, Any]] = []
            expectation_results: list[dict[str, Any]] = []
            event_expectation_results: list[dict[str, Any]] = []
            turn_summaries: list[dict[str, Any]] = []
            turn_outputs: list[dict[str, Any]] = []
            if args.case_timeout_s > 0:
                case_dir.mkdir(parents=True, exist_ok=True)
                worker_result_path = case_dir / "worker_result.case.json"
                if worker_result_path.exists():
                    worker_result_path.unlink()
                timeout_budget_s = timeout_s * max(1, len(turns))
                ctx = multiprocessing.get_context("spawn")
                process = ctx.Process(
                    target=_execute_case_worker,
                    args=(
                        str(worker_result_path),
                        env,
                        case_thread_key,
                        [turn.prompt for turn in turns],
                        args.model,
                    ),
                )
                process.start()
                deadline = time.monotonic() + timeout_budget_s
                while process.is_alive():
                    remaining = deadline - time.monotonic()
                    if remaining <= 0:
                        process.terminate()
                        process.join()
                        _force_remove_case_container(case_thread_key)
                        raise EvalCaseTimeoutError(
                            f"case execution timed out (timeout_s={timeout_s}, turns={len(turns)})"
                        )
                    process.join(min(1.0, remaining))
                if not worker_result_path.exists():
                    raise RuntimeError("case worker exited without writing a result")
                worker_result = json.loads(worker_result_path.read_text())
                if not worker_result.get("ok"):
                    err = worker_result.get("error", "unknown case worker failure")
                    tb = worker_result.get("traceback")
                    raise RuntimeError(f"{err}\n{tb}" if tb else err)
                turn_outputs = [
                    item for item in worker_result.get("turn_outputs", []) if isinstance(item, dict)
                ]
            else:
                if client is None:
                    raise RuntimeError("agent client is not initialized")
                for turn_idx, turn in enumerate(turns, start=1):
                    continue_session = turn_idx > 1
                    emitted_events: list[dict[str, Any]] = []

                    def _emit(event: dict[str, Any], sink: list[dict[str, Any]] = emitted_events) -> None:
                        if isinstance(event, dict):
                            sink.append(event)

                    response = client.execute(
                        case_thread_key,
                        turn.prompt,
                        harness="legal",
                        source="api",
                        model=args.model,
                        continue_session=continue_session,
                        emit=_emit,
                    )
                    if response.get("error"):
                        raise RuntimeError(f"agent returned error: {response['error']} (turn={turn_idx})")
                    turn_outputs.append(
                        {
                            "turn_index": turn_idx,
                            "continue_session": continue_session,
                            "response": dict(response),
                            "events": emitted_events,
                        }
                    )

            if len(turn_outputs) != len(turns):
                raise RuntimeError(
                    f"worker turn count mismatch: expected {len(turns)}, got {len(turn_outputs)}"
                )

            for turn_idx, turn in enumerate(turns, start=1):
                turn_payload = turn_outputs[turn_idx - 1]
                response = dict(turn_payload.get("response", {}))
                turn_events = [
                    item
                    for item in turn_payload.get("events", [])
                    if isinstance(item, dict)
                ]
                if response.get("error"):
                    raise RuntimeError(f"agent returned error: {response['error']} (turn={turn_idx})")

                turn_text = str(response.get("result", ""))
                final_response = response
                final_text = turn_text
                all_events.extend(turn_events)

                for exp in turn.expectations:
                    passed = _regex_match(turn_text, exp)
                    expectation_results.append(
                        {
                            "name": f"t{turn_idx}.{exp.name}",
                            "passed": passed,
                            "pattern": exp.pattern,
                        }
                    )
                expectation_results.append(
                    {
                        "name": f"t{turn_idx}.non_empty_output",
                        "passed": bool(turn_text.strip()),
                        "pattern": r"\S+",
                    }
                )
                if turn.max_chars is not None:
                    expectation_results.append(
                        {
                            "name": f"t{turn_idx}.max_chars",
                            "passed": len(turn_text) <= turn.max_chars,
                            "pattern": f"len(output)<= {turn.max_chars}",
                        }
                    )

                turn_event_blob = "\n".join(
                    json.dumps(item, sort_keys=True, default=str) for item in turn_events
                )
                for event_exp in turn.event_expectations:
                    event_result = _event_expectation_result(turn_event_blob, event_exp)
                    event_result["name"] = f"t{turn_idx}.{event_result['name']}"
                    event_expectation_results.append(event_result)

                turn_summaries.append(
                    {
                        "turn_index": turn_idx,
                        "continue_session": bool(turn_payload.get("continue_session", turn_idx > 1)),
                        "prompt_sha256": _sha256_text(turn.prompt),
                        "output_sha256": _sha256_text(turn_text),
                        "chars": len(turn_text),
                        "event_count": len(turn_events),
                    }
                )

            all_checks = expectation_results + event_expectation_results
            passed = all(item["passed"] for item in all_checks)
            if args.capture_case_files:
                _write_text(output_path, final_text + ("\n" if not final_text.endswith("\n") else ""))
                _write_text(response_path, json.dumps(final_response, indent=2, default=str))
                _write_text(events_path, json.dumps(all_events, indent=2, default=str))
            results.append(
                {
                    "case_id": case.case_id,
                    "category": case.category,
                    "tier": case.tier,
                    "passed": passed,
                    "duration_s": round(time.time() - started, 3),
                    "harness": final_response.get("harness"),
                    "engine": final_response.get("engine"),
                    "persona": final_response.get("persona"),
                    "turn_count": len(turns),
                    "event_count": len(all_events),
                    "chars": len(final_text),
                    "prompt_sha256": _sha256_text(prompt_blob),
                    "output_sha256": _sha256_text(final_text),
                    "expectations": expectation_results,
                    "event_expectations": event_expectation_results,
                    "turns": turn_summaries,
                    "preview": " ".join(final_text.split())[:700],
                    "prompt_path": str(prompt_path) if args.capture_case_files else None,
                    "output_path": str(output_path) if args.capture_case_files else None,
                    "response_path": str(response_path) if args.capture_case_files else None,
                    "events_path": str(events_path) if args.capture_case_files else None,
                }
            )
            if args.include_full_output_json:
                results[-1]["prompt_text"] = prompt_blob
                results[-1]["output_text"] = final_text
                results[-1]["events"] = all_events
            print(f"passed={passed} chars={len(final_text)} turns={len(turns)}")
        except Exception as exc:
            if args.capture_case_files:
                _write_text(error_path, f"{exc}\n")
            results.append(
                {
                    "case_id": case.case_id,
                    "category": case.category,
                    "tier": case.tier,
                    "passed": False,
                    "duration_s": round(time.time() - started, 3),
                    "error": str(exc),
                    "prompt_path": str(prompt_path) if args.capture_case_files else None,
                    "output_path": str(output_path) if args.capture_case_files else None,
                    "response_path": str(response_path) if args.capture_case_files else None,
                    "error_path": str(error_path) if args.capture_case_files else None,
                }
            )
            if args.include_full_output_json:
                results[-1]["prompt_text"] = prompt_blob
            print(f"failed with exception: {exc}")
        finally:
            if client is not None:
                with contextlib.suppress(Exception):
                    client.stop(case_thread_key)
            _force_remove_case_container(case_thread_key)

    if client is not None:
        with contextlib.suppress(Exception):
            client.stop(shared_run_key)

    if not args.skip_docx:
        print("\n=== DOCX fidelity checks ===")
        results.extend(_run_docx_fidelity_eval())
    if not args.skip_knowledge:
        print("\n=== Knowledge wiring checks ===")
        results.extend(_run_knowledge_wiring_eval())
    print("\n=== Legal mode guard checks ===")
    results.extend(_run_mode_guard_eval())

    total_cases = len(results)
    passed_cases = sum(1 for r in results if r.get("passed", False))
    failed_cases = total_cases - passed_cases
    pass_rate = (passed_cases / total_cases * 100.0) if total_cases else 0.0

    by_category: dict[str, dict[str, Any]] = {}
    for item in results:
        category = item.get("category", "unknown")
        stats = by_category.setdefault(category, {"passed": 0, "total": 0, "pass_rate": 0.0})
        stats["total"] += 1
        if item.get("passed", False):
            stats["passed"] += 1
    for stats in by_category.values():
        stats["pass_rate"] = (stats["passed"] / stats["total"] * 100.0) if stats["total"] else 0.0

    report = {
        "timestamp": timestamp,
        "model": args.model,
        "env_file": str(env_file) if env_file else None,
        "artifacts_root": str(artifacts_root) if args.capture_case_files else None,
        "total_cases": total_cases,
        "passed_cases": passed_cases,
        "failed_cases": failed_cases,
        "pass_rate": pass_rate,
        "by_category": by_category,
        "results": results,
    }

    json_path = out_dir / f"legal-eval-{timestamp}.json"
    md_path = out_dir / f"legal-eval-{timestamp}.md"
    json_path.write_text(json.dumps(report, indent=2))
    md_path.write_text(_render_markdown(report))

    print("\n=== Summary ===")
    print(f"pass_rate={pass_rate:.1f}% ({passed_cases}/{total_cases})")
    print(f"json={json_path}")
    print(f"md={md_path}")
    return 0 if failed_cases == 0 else 2


if __name__ == "__main__":
    raise SystemExit(main())
